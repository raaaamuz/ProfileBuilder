import os
import re
import base64
import anthropic
from io import BytesIO
from django.conf import settings
from django.http import FileResponse, HttpResponse
from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from rest_framework.decorators import api_view, permission_classes
from rest_framework.views import APIView
from django.shortcuts import get_object_or_404
from django.core.files.base import ContentFile
from .models import UserProfile, ProfileDesign, ProfileNote
from .serializers import UserProfileSerializer, ProfileDesignSerializer, ProfileNoteSerializer


def embed_profile_picture_in_html(html_content, profile):
    """Replace profile picture URL with base64 data URL for PDF generation"""
    if not profile.profile_picture or not profile.profile_picture.name:
        return html_content

    try:
        # Read the profile picture and convert to base64
        profile.profile_picture.open('rb')
        image_data = profile.profile_picture.read()
        profile.profile_picture.close()

        # Determine mime type
        filename = profile.profile_picture.name.lower()
        if filename.endswith('.png'):
            mime_type = 'image/png'
        elif filename.endswith('.gif'):
            mime_type = 'image/gif'
        else:
            mime_type = 'image/jpeg'

        # Create base64 data URL
        base64_data = base64.b64encode(image_data).decode('utf-8')
        data_url = f'data:{mime_type};base64,{base64_data}'

        # Replace image src URLs that contain the profile picture path
        # Match various URL patterns for the profile picture
        picture_filename = os.path.basename(profile.profile_picture.name)
        patterns = [
            rf'src="[^"]*{re.escape(picture_filename)}[^"]*"',
            rf"src='[^']*{re.escape(picture_filename)}[^']*'",
            r'src="[^"]*profiles/[^"]*"',
            r"src='[^']*profiles/[^']*'",
        ]

        for pattern in patterns:
            html_content = re.sub(pattern, f'src="{data_url}"', html_content)

        return html_content
    except Exception as e:
        print(f"Error embedding profile picture: {e}")
        return html_content

# Configure Claude API
import os
CLAUDE_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
claude_client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

class UserProfileViewSet(viewsets.ModelViewSet):
    """
    Handles profile creation and updates.
    If a profile already exists for the user, a POST request will update it.
    """
    serializer_class = UserProfileSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return UserProfile.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    def create(self, request, *args, **kwargs):
        # Check if profile exists; if so, update it. Otherwise create new.
        user_profile, created = UserProfile.objects.get_or_create(user=request.user)
        serializer = self.get_serializer(user_profile, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data)

# Example API view to fetch any user's profile using their username.
@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def get_user_profile(request, username):
    """Allow home content to fetch user profile details"""
    profile = get_object_or_404(UserProfile, user__username=username)
    serializer = UserProfileSerializer(profile)
    return Response(serializer.data)


from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework import permissions
from django.shortcuts import get_object_or_404
from .models import UserProfile

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def get_contact_info(request):
    """
    API endpoint to retrieve the authenticated user's contact information.
    """
    # Get the UserProfile for the authenticated user.
    profile = get_object_or_404(UserProfile, user=request.user)
    
    # Collect contact information. Adjust field names as necessary.
    contact_info = {
        "email": request.user.email,
        "phone": getattr(profile, 'phone', None),
        "address": getattr(profile, 'address', None),
    }
    
    return Response(contact_info)


# ========== Resume Templates ==========

RESUME_TEMPLATES = {
    'professional': {
        'name': 'Professional',
        'description': 'Clean two-column layout with dark blue header',
        'colors': {'primary': '#1a365d', 'secondary': '#3182ce', 'accent': '#90cdf4', 'bg': '#f7fafc'},
    },
    'modern': {
        'name': 'Modern',
        'description': 'Bold header with purple accent colors',
        'colors': {'primary': '#5b21b6', 'secondary': '#7c3aed', 'accent': '#c4b5fd', 'bg': '#faf5ff'},
    },
    'minimal': {
        'name': 'Minimal',
        'description': 'Simple, elegant single-column design',
        'colors': {'primary': '#1f2937', 'secondary': '#4b5563', 'accent': '#9ca3af', 'bg': '#ffffff'},
    },
    'executive': {
        'name': 'Executive',
        'description': 'Sophisticated dark theme for senior professionals',
        'colors': {'primary': '#111827', 'secondary': '#d97706', 'accent': '#fbbf24', 'bg': '#1f2937'},
    },
    'creative': {
        'name': 'Creative',
        'description': 'Colorful sidebar with modern typography',
        'colors': {'primary': '#0d9488', 'secondary': '#14b8a6', 'accent': '#5eead4', 'bg': '#f0fdfa'},
    },
    'tech': {
        'name': 'Tech',
        'description': 'Developer-focused with code-style elements',
        'colors': {'primary': '#0f172a', 'secondary': '#22c55e', 'accent': '#86efac', 'bg': '#020617'},
    },
    'elegant': {
        'name': 'Elegant',
        'description': 'Luxurious purple gradient design',
        'colors': {'primary': '#1e1b4b', 'secondary': '#c084fc', 'accent': '#e9d5ff', 'bg': '#faf5ff'},
    },
    'corporate': {
        'name': 'Corporate',
        'description': 'Professional business layout',
        'colors': {'primary': '#0c4a6e', 'secondary': '#0284c7', 'accent': '#7dd3fc', 'bg': '#f0f9ff'},
    },
    'bold': {
        'name': 'Bold',
        'description': 'Eye-catching red and orange design',
        'colors': {'primary': '#dc2626', 'secondary': '#f87171', 'accent': '#fecaca', 'bg': '#fef2f2'},
    },
    'nature': {
        'name': 'Nature',
        'description': 'Fresh green organic theme',
        'colors': {'primary': '#166534', 'secondary': '#22c55e', 'accent': '#bbf7d0', 'bg': '#f0fdf4'},
    },
    'midnight': {
        'name': 'Midnight',
        'description': 'Deep dark blue with indigo accents',
        'colors': {'primary': '#020617', 'secondary': '#6366f1', 'accent': '#a5b4fc', 'bg': '#0f172a'},
    },
    'sunset': {
        'name': 'Sunset',
        'description': 'Warm orange and amber tones',
        'colors': {'primary': '#9a3412', 'secondary': '#f97316', 'accent': '#fed7aa', 'bg': '#fff7ed'},
    },
    'swiss': {
        'name': 'Swiss',
        'description': 'Clean Swiss design with red accent',
        'colors': {'primary': '#000000', 'secondary': '#dc2626', 'accent': '#fca5a5', 'bg': '#ffffff'},
    },
    'nordic': {
        'name': 'Nordic',
        'description': 'Scandinavian minimal style',
        'colors': {'primary': '#374151', 'secondary': '#6b7280', 'accent': '#d1d5db', 'bg': '#f9fafb'},
    },
    'paper': {
        'name': 'Paper',
        'description': 'Clean paper-like traditional design',
        'colors': {'primary': '#292524', 'secondary': '#78716c', 'accent': '#d6d3d1', 'bg': '#fafaf9'},
    },
    'mono': {
        'name': 'Mono',
        'description': 'Monochromatic minimal black and white',
        'colors': {'primary': '#18181b', 'secondary': '#52525b', 'accent': '#a1a1aa', 'bg': '#ffffff'},
    },
    'classic': {
        'name': 'Classic',
        'description': 'Traditional clean resume layout',
        'colors': {'primary': '#1e3a5f', 'secondary': '#3b82f6', 'accent': '#bfdbfe', 'bg': '#ffffff'},
    },
}

# Complete template CSS that can restyle the same HTML structure
RESUME_TEMPLATE_CSS = {
    'professional': '''
        :root { --primary: #1a365d; --secondary: #3182ce; --accent: #90cdf4; --bg: #f7fafc; --text: #1a202c; --light-text: #718096; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: linear-gradient(135deg, #1a365d 0%, #2c5282 100%); color: white; padding: 25px; }
        .header h1 { color: white; }
        .header .contact { color: rgba(255,255,255,0.9); }
        .sidebar { background: #f1f5f9; }
        .section-title { color: var(--primary); border-bottom: 2px solid var(--secondary); }
        .skill-tag { background: var(--secondary); color: white; }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
        .job-date { color: var(--secondary); }
    ''',
    'modern': '''
        :root { --primary: #5b21b6; --secondary: #7c3aed; --accent: #c4b5fd; --bg: #faf5ff; --text: #1f2937; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: linear-gradient(135deg, #7c3aed 0%, #5b21b6 100%); color: white; padding: 25px; }
        .header h1 { color: white; }
        .sidebar { background: #f5f3ff; }
        .section-title { color: var(--primary); border-bottom: 2px solid var(--secondary); }
        .skill-tag { background: #ede9fe; color: var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
        .job-date { color: var(--secondary); }
    ''',
    'minimal': '''
        :root { --primary: #1f2937; --secondary: #6b7280; --text: #1f2937; }
        body { background: white; }
        .resume { background: white; }
        .header { background: white; color: var(--primary); border-bottom: 2px solid #e5e7eb; padding: 20px 0; }
        .header h1 { color: var(--primary); }
        .sidebar { background: white; border-right: 1px solid #e5e7eb; }
        .section-title { color: var(--primary); border-bottom: 1px solid #e5e7eb; }
        .skill-tag { background: #f3f4f6; color: var(--primary); border: 1px solid #e5e7eb; }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
    ''',
    'executive': '''
        :root { --primary: #fbbf24; --secondary: #d97706; --bg: #111827; --text: #e5e7eb; }
        body { background: var(--bg); color: var(--text); }
        .resume { background: #1f2937; color: var(--text); }
        .header { background: #111827; color: white; border-bottom: 2px solid var(--secondary); padding: 25px; }
        .header h1 { color: var(--primary); }
        .header .contact { color: #9ca3af; }
        .sidebar { background: #111827; color: #e5e7eb; }
        .section-title { color: var(--primary); border-bottom: 1px solid var(--secondary); }
        .skill-tag { background: #374151; color: var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
        .job-date { color: #9ca3af; }
        .job-desc, p, li { color: #d1d5db; }
    ''',
    'creative': '''
        :root { --primary: #0d9488; --secondary: #14b8a6; --accent: #5eead4; --bg: #f0fdfa; --text: #134e4a; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: linear-gradient(135deg, #0d9488 0%, #06b6d4 100%); color: white; padding: 25px; }
        .header h1 { color: white; }
        .sidebar { background: linear-gradient(180deg, #0d9488 0%, #14b8a6 100%); color: white; }
        .sidebar .section-title { color: white; border-bottom-color: rgba(255,255,255,0.3); }
        .sidebar .skill-tag { background: rgba(255,255,255,0.2); color: white; }
        .section-title { color: var(--primary); border-bottom: 2px solid var(--secondary); }
        .skill-tag { background: var(--secondary); color: white; }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
        .job-date { color: var(--secondary); }
    ''',
    'tech': '''
        :root { --primary: #22c55e; --secondary: #16a34a; --bg: #0f172a; --text: #e2e8f0; }
        body { background: var(--bg); color: var(--text); }
        .resume { background: #0f172a; color: var(--text); border: 1px solid #1e293b; }
        .header { background: #020617; color: white; padding: 25px; border-bottom: 1px solid #1e293b; }
        .header h1 { color: var(--primary); font-family: 'Courier New', monospace; }
        .header .contact { color: #94a3b8; }
        .sidebar { background: #1e293b; color: #e2e8f0; }
        .section-title { color: var(--primary); font-family: 'Courier New', monospace; border-bottom: 1px solid #1e293b; }
        .skill-tag { background: #1e293b; color: var(--primary); border: 1px solid var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: #94a3b8; }
        .job-date { color: var(--primary); }
        .job-desc, p, li { color: #cbd5e1; }
    ''',
    'elegant': '''
        :root { --primary: #c084fc; --secondary: #a855f7; --bg: #1e1b4b; --text: #e9d5ff; }
        body { background: var(--bg); color: var(--text); }
        .resume { background: linear-gradient(180deg, #1e1b4b 0%, #312e81 100%); color: var(--text); }
        .header { background: transparent; color: white; padding: 25px; border-bottom: 1px solid rgba(192,132,252,0.3); }
        .header h1 { color: var(--primary); }
        .header .contact { color: #c4b5fd; }
        .sidebar { background: rgba(139,92,246,0.1); color: #e9d5ff; }
        .section-title { color: var(--primary); border-bottom: 1px solid rgba(192,132,252,0.3); }
        .skill-tag { background: rgba(139,92,246,0.2); color: var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: #c4b5fd; }
        .job-date { color: #a78bfa; }
        .job-desc, p, li { color: #ddd6fe; }
    ''',
    'corporate': '''
        :root { --primary: #0c4a6e; --secondary: #0284c7; --accent: #7dd3fc; --bg: #f0f9ff; --text: #0c4a6e; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: linear-gradient(135deg, #0c4a6e 0%, #0369a1 100%); color: white; padding: 25px; }
        .header h1 { color: white; }
        .header .contact { color: rgba(255,255,255,0.9); }
        .sidebar { background: #f0f9ff; border-right: 2px solid #bae6fd; }
        .section-title { color: var(--primary); border-bottom: 2px solid var(--secondary); }
        .skill-tag { background: #e0f2fe; color: var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
        .job-date { color: var(--secondary); }
    ''',
    'bold': '''
        :root { --primary: #dc2626; --secondary: #f97316; --accent: #fecaca; --bg: #fef2f2; --text: #1f2937; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: linear-gradient(135deg, #dc2626 0%, #f97316 100%); color: white; padding: 25px; }
        .header h1 { color: white; }
        .header .contact { color: rgba(255,255,255,0.9); }
        .sidebar { background: #fef2f2; }
        .section-title { color: var(--primary); border-bottom: 2px solid var(--secondary); }
        .skill-tag { background: linear-gradient(135deg, #fee2e2, #ffedd5); color: var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
        .job-date { color: var(--secondary); }
    ''',
    'nature': '''
        :root { --primary: #166534; --secondary: #22c55e; --accent: #bbf7d0; --bg: #f0fdf4; --text: #14532d; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: linear-gradient(135deg, #166534 0%, #15803d 100%); color: white; padding: 25px; }
        .header h1 { color: white; }
        .header .contact { color: rgba(255,255,255,0.9); }
        .sidebar { background: linear-gradient(180deg, #f0fdf4 0%, #dcfce7 100%); }
        .section-title { color: var(--primary); border-bottom: 2px solid var(--secondary); }
        .skill-tag { background: #dcfce7; color: var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
        .job-date { color: var(--secondary); }
    ''',
    'midnight': '''
        :root { --primary: #6366f1; --secondary: #818cf8; --bg: #020617; --text: #e2e8f0; }
        body { background: var(--bg); color: var(--text); }
        .resume { background: #0f172a; color: var(--text); }
        .header { background: #020617; color: white; padding: 25px; border-bottom: 1px solid rgba(99,102,241,0.3); }
        .header h1 { color: var(--primary); }
        .header .contact { color: #a5b4fc; }
        .sidebar { background: #1e1b4b; color: #e2e8f0; }
        .section-title { color: var(--primary); border-bottom: 1px solid rgba(99,102,241,0.3); }
        .skill-tag { background: rgba(99,102,241,0.2); color: var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: #a5b4fc; }
        .job-date { color: #818cf8; }
        .job-desc, p, li { color: #cbd5e1; }
    ''',
    'sunset': '''
        :root { --primary: #9a3412; --secondary: #f97316; --accent: #fed7aa; --bg: #fff7ed; --text: #1f2937; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: linear-gradient(135deg, #ea580c 0%, #f59e0b 100%); color: white; padding: 25px; }
        .header h1 { color: white; }
        .header .contact { color: rgba(255,255,255,0.9); }
        .sidebar { background: linear-gradient(180deg, #fff7ed 0%, #ffedd5 100%); }
        .section-title { color: var(--primary); border-bottom: 2px solid var(--secondary); }
        .skill-tag { background: #ffedd5; color: var(--primary); }
        .job-title { color: var(--primary); }
        .job-company { color: var(--secondary); }
        .job-date { color: var(--secondary); }
    ''',
    'swiss': '''
        :root { --primary: #000000; --secondary: #dc2626; --accent: #fca5a5; --bg: #ffffff; --text: #000000; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: white; color: black; padding: 25px; border-left: 6px solid #dc2626; }
        .header h1 { color: black; font-weight: 800; }
        .header .contact { color: #4b5563; }
        .sidebar { background: white; border-right: 1px solid #e5e7eb; }
        .section-title { color: black; border-bottom: 2px solid #dc2626; font-weight: 700; }
        .skill-tag { background: #fee2e2; color: #dc2626; }
        .job-title { color: black; font-weight: 600; }
        .job-company { color: #dc2626; }
        .job-date { color: #6b7280; }
    ''',
    'nordic': '''
        :root { --primary: #374151; --secondary: #6b7280; --accent: #d1d5db; --bg: #f9fafb; --text: #374151; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: #f9fafb; color: #374151; padding: 25px; text-align: center; border-bottom: 1px solid #e5e7eb; }
        .header h1 { color: #1f2937; }
        .header .contact { color: #6b7280; }
        .sidebar { background: #f9fafb; }
        .section-title { color: #374151; border-bottom: 1px solid #d1d5db; }
        .skill-tag { background: #f3f4f6; color: #374151; border: 1px solid #e5e7eb; }
        .job-title { color: #1f2937; }
        .job-company { color: #6b7280; }
        .job-date { color: #9ca3af; }
    ''',
    'paper': '''
        :root { --primary: #292524; --secondary: #78716c; --accent: #d6d3d1; --bg: #fafaf9; --text: #292524; }
        body { background: var(--bg); }
        .resume { background: #fafaf9; }
        .header { background: #fafaf9; color: #292524; padding: 25px; border-bottom: 3px solid #292524; }
        .header h1 { color: #292524; }
        .header .contact { color: #78716c; }
        .sidebar { background: #fafaf9; border-right: 1px solid #d6d3d1; }
        .section-title { color: #292524; border-bottom: 2px solid #a8a29e; }
        .skill-tag { background: #e7e5e4; color: #292524; }
        .job-title { color: #292524; }
        .job-company { color: #78716c; }
        .job-date { color: #a8a29e; }
    ''',
    'mono': '''
        :root { --primary: #18181b; --secondary: #52525b; --accent: #a1a1aa; --bg: #ffffff; --text: #18181b; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: #18181b; color: white; padding: 25px; }
        .header h1 { color: white; }
        .header .contact { color: #a1a1aa; }
        .sidebar { background: #fafafa; }
        .section-title { color: #18181b; border-bottom: 2px solid #18181b; }
        .skill-tag { background: #f4f4f5; color: #18181b; border: 1px solid #e4e4e7; }
        .job-title { color: #18181b; }
        .job-company { color: #52525b; }
        .job-date { color: #71717a; }
    ''',
    'classic': '''
        :root { --primary: #1e3a5f; --secondary: #3b82f6; --accent: #bfdbfe; --bg: #ffffff; --text: #1f2937; }
        body { background: var(--bg); }
        .resume { background: white; }
        .header { background: white; color: #1e3a5f; padding: 25px; text-align: center; border-bottom: 2px solid #1e3a5f; }
        .header h1 { color: #1e3a5f; }
        .header .contact { color: #6b7280; }
        .sidebar { background: white; }
        .section-title { color: #1e3a5f; border-bottom: 1px solid #3b82f6; }
        .skill-tag { background: #eff6ff; color: #1e3a5f; }
        .job-title { color: #1e3a5f; }
        .job-company { color: #3b82f6; }
        .job-date { color: #6b7280; }
    ''',
}


# ========== AI Resume API Views ==========

@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def resume_templates(request):
    """Get available resume templates"""
    templates = []
    for key, value in RESUME_TEMPLATES.items():
        templates.append({
            'id': key,
            'name': value['name'],
            'description': value['description'],
            'colors': value['colors'],
        })
    return Response({'templates': templates})


@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def resume_template_css(request, template_id):
    """Get CSS for a specific template to apply dynamically"""
    css = RESUME_TEMPLATE_CSS.get(template_id, RESUME_TEMPLATE_CSS['professional'])
    return Response({'css': css, 'template': template_id})


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def resume_status(request):
    """Check if user has an approved AI resume"""
    profile = get_object_or_404(UserProfile, user=request.user)
    has_resume = bool(profile.ai_resume and profile.ai_resume.name)
    return Response({
        'has_resume': has_resume,
        'resume_url': profile.ai_resume.url if has_resume else None
    })


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def resume_generate(request):
    """Generate resume using Claude AI with multiple template options"""
    from career.models import CareerEntry, Skill
    from education.models import EducationEntry
    from achievements.models import Award, Achievement

    user = request.user
    profile = get_object_or_404(UserProfile, user=user)

    # Get job description for ATS optimization
    job_description = request.data.get('job_description', '').strip()

    # Get selected template (default to professional)
    template_id = request.data.get('template', 'professional')
    template = RESUME_TEMPLATES.get(template_id, RESUME_TEMPLATES['professional'])
    colors = template['colors']

    # Get section visibility settings
    sections = request.data.get('sections', {})
    include_summary = sections.get('summary', True)
    include_experience = sections.get('experience', True)
    include_education = sections.get('education', True)
    include_skills = sections.get('skills', True)
    include_awards = sections.get('awards', True)
    include_achievements = sections.get('achievements', True)

    # Get section order (default order if not provided)
    section_order = request.data.get('section_order', ['summary', 'experience', 'education', 'skills', 'awards', 'achievements'])

    # Get section styles
    section_styles = request.data.get('section_styles', {})

    # Get column placement for two-column layouts
    sidebar_sections = request.data.get('sidebar_sections', ['skills', 'education'])
    main_sections = request.data.get('main_sections', ['summary', 'experience', 'awards', 'achievements'])
    print(f"[DEBUG] Received section_styles: {section_styles}")

    # Gather user data
    careers = CareerEntry.objects.filter(user=user).order_by('-year')
    educations = EducationEntry.objects.filter(user=user).order_by('-year')
    skills = Skill.objects.filter(user=user)
    awards = Award.objects.filter(user=user).order_by('-year')
    achievements = Achievement.objects.filter(user=user).order_by('-year')

    # Build career section with details
    # Check if achievement-focused style is selected
    experience_style = section_styles.get('experience', 'detailed')

    career_text = ""
    for c in careers:
        if experience_style == 'achievements':
            # For achievement style, instruct AI to convert to bullet points
            career_text += f"""
JOB (CONVERT TO ACHIEVEMENT BULLETS - NO PARAGRAPHS):
- Title: {c.title}
- Company: {c.company}
- Period: {c.year}
- Raw Description (MUST CONVERT TO 3-5 BULLET POINTS WITH METRICS): {c.description or 'N/A'}
[INSTRUCTION: Transform the above description into 3-5 achievement bullet points. Each bullet MUST:
 1. Start with action verb (Led, Developed, Increased, Managed, Built, etc.)
 2. Include specific metrics/numbers (%, $, team size, volume)
 3. Be ONE LINE, not a paragraph
 Example: "• Led team of 12 analysts, improving report delivery time by 40%"]
"""
        else:
            career_text += f"""
JOB:
- Title: {c.title}
- Company: {c.company}
- Period: {c.year}
- Description: {c.description or 'N/A'}
"""

    # Build education section
    education_text = ""
    for e in educations:
        education_text += f"- {e.degree} from {e.university} ({e.year})\n"
        if e.description:
            education_text += f"  Details: {e.description}\n"

    # Build skills section with proficiency (neutral format - no stars)
    skills_list = []
    for s in skills:
        # Send skill name with proficiency percentage for flexible styling
        proficiency_pct = s.proficiency * 20  # Convert 1-5 to percentage
        skills_list.append(f"{s.name}: {proficiency_pct}%")
    skills_text = ", ".join(skills_list) if skills_list else "Not specified"

    # Build awards section
    awards_text = ""
    for a in awards:
        awards_text += f"- {a.title} from {a.organization} ({a.year})\n"

    # Build achievements section
    achievements_text = ""
    for a in achievements:
        achievements_text += f"- {a.title} ({a.year}): {a.description}\n"

    # Get full name
    full_name = f"{user.first_name} {user.last_name}".strip() or user.username

    # Get profile picture URL
    profile_picture_url = ""
    if profile.profile_picture and profile.profile_picture.name:
        profile_picture_url = request.build_absolute_uri(profile.profile_picture.url)

    # Build section order instruction for AI
    section_labels = {
        'summary': 'Professional Summary',
        'experience': 'Work Experience',
        'education': 'Education',
        'skills': 'Skills',
        'awards': 'Awards',
        'achievements': 'Achievements'
    }
    ordered_sections = [section_labels.get(s, s.title()) for s in section_order if sections.get(s, True)]

    # Build column placement instruction
    sidebar_section_names = [section_labels.get(s, s.title()) for s in sidebar_sections if sections.get(s, True)]
    main_section_names = [section_labels.get(s, s.title()) for s in main_sections if sections.get(s, True)]

    column_placement = ""
    if sidebar_section_names or main_section_names:
        column_placement = f"""
COLUMN PLACEMENT (for two-column layouts):
- LEFT SIDEBAR should contain: {', '.join(sidebar_section_names) if sidebar_section_names else 'None'}
- RIGHT MAIN CONTENT should contain: {', '.join(main_section_names) if main_section_names else 'None'}
"""

    section_order_instruction = f"""
IMPORTANT - SECTION ORDER AND PLACEMENT:
Generate the resume sections in EXACTLY this order:
{chr(10).join(f'{i+1}. {s}' for i, s in enumerate(ordered_sections))}
{column_placement}
Do NOT include sections that are not in this list. Follow this order precisely.
"""

    # Build section styles instruction
    style_descriptions = {
        'skills': {
            'chips': 'Display skills as colorful rounded badge/chip tags in a flex-wrap layout. NO stars or proficiency indicators - just skill names in colored pill-shaped tags',
            'bars': 'Show each skill with a horizontal progress bar indicating proficiency level (percentage width)',
            'grid': 'Organize skills in a clean 3-column grid layout with skill names only, no stars',
            'list': 'Simple bullet-point list of skill names only, no proficiency stars',
        },
        'education': {
            'cards': 'Each education entry in a boxed card with shadow',
            'timeline': 'Vertical timeline with dots and connecting lines',
            'compact': 'Minimal space-efficient listing',
        },
        'summary': {
            'paragraph': 'Full paragraph text block',
            'highlights': 'Key points with bullet icons or checkmarks',
            'quote': 'Styled as an elegant quote with decorative borders',
        },
        'experience': {
            'detailed': 'Full job descriptions with bullet points for responsibilities',
            'timeline': 'Visual timeline with dates on one side, content on other',
            'achievements': 'Focus on achievements with bullet points, metrics highlighted',
            'compact': 'Space-efficient with essential info only',
        },
        'awards': {
            'cards': 'Award cards with trophy/ribbon icons',
            'list': 'Clean simple listing',
            'badges': 'Badge or ribbon-style decorative elements',
        },
        'achievements': {
            'bullets': 'Clear bullet point list',
            'numbered': 'Numbered list showing ranking',
            'icons': 'Each achievement with a relevant icon',
        },
    }

    # HTML patterns for each skill style
    skills_html_patterns = {
        'chips': '''SKILLS HTML - USE EXACTLY THIS PATTERN (colorful pill badges):
<div style="display: flex; flex-wrap: wrap; gap: 8px;">
  <span style="background: #3b82f6; color: white; padding: 4px 12px; border-radius: 20px; font-size: 12px;">SkillName</span>
  <!-- Repeat for each skill, use different colors: #22c55e, #8b5cf6, #f59e0b, #ef4444, #06b6d4 -->
</div>
DO NOT show percentages or proficiency levels. Just skill names in colored pills.''',

        'bars': '''SKILLS HTML - USE EXACTLY THIS PATTERN (progress bars with percentage):
For EACH skill, create this structure:
<div style="margin-bottom: 10px;">
  <div style="display: flex; justify-content: space-between; font-size: 11px; margin-bottom: 3px;">
    <span style="font-weight: 500;">SkillName</span>
    <span style="color: #6b7280;">XX%</span>
  </div>
  <div style="height: 6px; background: #e5e7eb; border-radius: 3px; overflow: hidden;">
    <div style="width: XX%; height: 100%; background: #3b82f6; border-radius: 3px;"></div>
  </div>
</div>
Replace XX with the actual percentage from skills data. Each skill MUST have name, percentage, and visual bar.''',

        'grid': '''SKILLS HTML - USE EXACTLY THIS PATTERN (3-column grid):
<div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 6px;">
  <div style="background: #f3f4f6; padding: 6px 8px; border-radius: 4px; text-align: center; font-size: 11px;">SkillName</div>
  <!-- Repeat for each skill -->
</div>
DO NOT show percentages. Just skill names in grid boxes.''',

        'list': '''SKILLS HTML - USE EXACTLY THIS PATTERN (bullet list):
<ul style="list-style: disc; padding-left: 20px; margin: 0;">
  <li style="font-size: 11px; margin-bottom: 2px;">SkillName</li>
</ul>
DO NOT show percentages. Just skill names as bullet points.'''
    }

    # HTML patterns for summary styles
    summary_html_patterns = {
        'highlights': '''SUMMARY HTML - USE EXACTLY THIS PATTERN (bullet highlights):
<div style="padding: 12px; background: #f8fafc; border-radius: 8px; border-left: 4px solid #3b82f6;">
  <div style="display: flex; flex-wrap: wrap; gap: 8px;">
    <span style="display: inline-flex; align-items: center; gap: 4px; font-size: 11px;">
      <span style="color: #22c55e;">✓</span> Highlight point 1
    </span>
    <span style="display: inline-flex; align-items: center; gap: 4px; font-size: 11px;">
      <span style="color: #22c55e;">✓</span> Highlight point 2
    </span>
  </div>
</div>
Extract 4-6 KEY highlights from the bio/summary. Each highlight should be a short impactful phrase (3-6 words) with a green checkmark.''',

        'quote': '''SUMMARY HTML - USE EXACTLY THIS PATTERN (elegant quote):
<div style="border-left: 4px solid #8b5cf6; padding: 15px 20px; background: linear-gradient(to right, #f5f3ff, transparent); margin: 10px 0;">
  <p style="font-style: italic; font-size: 12px; color: #4b5563; line-height: 1.6; margin: 0;">
    "The summary text here as an elegant quote..."
  </p>
</div>
Format the summary as an elegant italicized quote with decorative left border.''',

        'paragraph': '''SUMMARY: Use standard paragraph format with the summary text.'''
    }

    # HTML patterns for experience/work experience styles
    experience_html_patterns = {
        'achievements': '''WORK EXPERIENCE HTML - USE ACHIEVEMENT-FOCUSED FORMAT:
For EACH job entry, structure it as:
<div style="margin-bottom: 20px;">
  <div style="display: flex; justify-content: space-between; align-items: baseline;">
    <div>
      <h4 style="font-weight: 600; font-size: 14px; margin: 0;">Job Title</h4>
      <p style="color: #3b82f6; font-size: 12px; margin: 2px 0;">Company Name</p>
    </div>
    <span style="font-size: 11px; color: #6b7280;">Date Range</span>
  </div>
  <ul style="margin: 8px 0 0 0; padding-left: 16px; font-size: 11px; line-height: 1.5;">
    <li style="margin-bottom: 4px;"><strong>Achievement:</strong> Specific measurable result (e.g., "Increased sales by 40%")</li>
    <li style="margin-bottom: 4px;"><strong>Achievement:</strong> Another quantified accomplishment with metrics</li>
    <li style="margin-bottom: 4px;"><strong>Achievement:</strong> Key impact with numbers or percentages</li>
  </ul>
</div>

CRITICAL FOR ACHIEVEMENTS STYLE:
- Convert job descriptions into 3-5 BULLET POINTS per job
- Each bullet MUST start with an action verb (Led, Developed, Increased, Reduced, etc.)
- Include METRICS/NUMBERS wherever possible (percentages, dollar amounts, team sizes)
- Focus on IMPACT and RESULTS, not just responsibilities
- NO long paragraphs - ONLY concise achievement bullets
- Example transformations:
  * "Managed team" → "Led cross-functional team of 8 engineers, delivering project 2 weeks ahead of schedule"
  * "Handled data" → "Processed 10M+ daily records, reducing data latency by 60%"
  * "Worked on sales" → "Increased quarterly revenue by $500K through strategic client outreach"
''',

        'detailed': '''WORK EXPERIENCE HTML - USE DETAILED FORMAT:
For EACH job entry:
<div style="margin-bottom: 20px;">
  <h4 style="font-weight: 600; font-size: 14px; margin: 0;">Job Title</h4>
  <p style="color: #6b7280; font-size: 12px;">Company Name | Date Range</p>
  <p style="font-size: 11px; line-height: 1.6; margin-top: 8px; color: #374151;">
    Full description paragraph with responsibilities and achievements...
  </p>
</div>
Use the full job description as provided, formatted as a readable paragraph.''',

        'timeline': '''WORK EXPERIENCE HTML - USE TIMELINE FORMAT:
<div style="position: relative; padding-left: 20px; border-left: 2px solid #e5e7eb;">
  <div style="position: relative; margin-bottom: 20px;">
    <div style="position: absolute; left: -26px; width: 10px; height: 10px; background: #3b82f6; border-radius: 50%;"></div>
    <span style="font-size: 11px; color: #6b7280;">Date Range</span>
    <h4 style="font-weight: 600; font-size: 14px; margin: 4px 0;">Job Title</h4>
    <p style="color: #3b82f6; font-size: 12px;">Company Name</p>
    <p style="font-size: 11px; margin-top: 6px;">Brief description...</p>
  </div>
</div>
Create a visual timeline with dates, connecting line, and dots for each position.''',

        'compact': '''WORK EXPERIENCE HTML - USE COMPACT FORMAT:
<div style="margin-bottom: 12px;">
  <div style="display: flex; justify-content: space-between;">
    <strong style="font-size: 12px;">Job Title</strong>
    <span style="font-size: 11px; color: #6b7280;">Date</span>
  </div>
  <p style="font-size: 11px; color: #6b7280; margin: 2px 0;">Company Name</p>
</div>
Space-efficient listing with essential info only. Brief or no description.'''
    }

    # HTML patterns for education styles
    education_html_patterns = {
        'timeline': '''EDUCATION HTML - USE TIMELINE FORMAT:
<div style="position: relative; padding-left: 20px; border-left: 2px solid #e5e7eb;">
  <div style="position: relative; margin-bottom: 16px;">
    <div style="position: absolute; left: -26px; width: 10px; height: 10px; background: #dc2626; border-radius: 50%;"></div>
    <h4 style="font-weight: 600; font-size: 13px; margin: 0;">Degree Name</h4>
    <p style="font-size: 11px; color: #6b7280; margin: 2px 0;">University Name • Year</p>
  </div>
</div>
Create a visual timeline with colored dots and connecting line.''',

        'cards': '''EDUCATION HTML - USE CARD FORMAT:
<div style="background: #f8fafc; border-radius: 8px; padding: 12px; margin-bottom: 10px; border: 1px solid #e5e7eb;">
  <h4 style="font-weight: 600; font-size: 13px; margin: 0;">Degree Name</h4>
  <p style="font-size: 11px; color: #6b7280; margin: 4px 0;">University Name</p>
  <span style="font-size: 10px; color: #9ca3af;">Year</span>
</div>
Each education entry in a card with subtle background and border.''',

        'compact': '''EDUCATION HTML - USE COMPACT FORMAT:
<div style="margin-bottom: 8px;">
  <strong style="font-size: 12px;">Degree Name</strong>
  <p style="font-size: 11px; color: #6b7280; margin: 0;">University • Year</p>
</div>
Minimal, space-efficient format.'''
    }

    # HTML patterns for awards styles
    awards_html_patterns = {
        'cards': '''AWARDS HTML - USE CARD FORMAT:
<div style="display: flex; align-items: center; gap: 10px; background: #fffbeb; border-radius: 8px; padding: 10px; margin-bottom: 8px; border: 1px solid #fcd34d;">
  <span style="font-size: 20px;">🏆</span>
  <div>
    <h4 style="font-weight: 600; font-size: 12px; margin: 0;">Award Title</h4>
    <p style="font-size: 11px; color: #6b7280; margin: 0;">Organization • Year</p>
  </div>
</div>
Award cards with trophy icon, warm background color.''',

        'list': '''AWARDS HTML - USE LIST FORMAT:
<ul style="list-style: none; padding: 0; margin: 0;">
  <li style="font-size: 11px; margin-bottom: 6px; padding-left: 16px; position: relative;">
    <span style="position: absolute; left: 0;">•</span>
    <strong>Award Title</strong> - Organization (Year)
  </li>
</ul>
Clean bullet list format.''',

        'badges': '''AWARDS HTML - USE BADGE FORMAT:
<div style="display: flex; flex-wrap: wrap; gap: 8px;">
  <div style="display: inline-flex; align-items: center; gap: 6px; background: linear-gradient(135deg, #fcd34d, #f59e0b); color: #78350f; padding: 6px 12px; border-radius: 20px; font-size: 11px; font-weight: 500;">
    🏅 Award Title (Year)
  </div>
</div>
Decorative badge-style ribbons with gradient background.'''
    }

    section_styles_instruction = ""
    skills_html_instruction = ""
    summary_html_instruction = ""
    experience_html_instruction = ""
    education_html_instruction = ""
    awards_html_instruction = ""

    for section_key, style_id in section_styles.items():
        if section_key in style_descriptions and style_id in style_descriptions[section_key]:
            section_styles_instruction += f"\n- {section_key.upper()}: Use '{style_id}' style - {style_descriptions[section_key][style_id]}"

            # Add specific HTML pattern for skills
            if section_key == 'skills' and style_id in skills_html_patterns:
                skills_html_instruction = skills_html_patterns[style_id]

            # Add specific HTML pattern for summary
            if section_key == 'summary' and style_id in summary_html_patterns:
                summary_html_instruction = summary_html_patterns[style_id]

            # Add specific HTML pattern for experience/work experience
            if section_key == 'experience' and style_id in experience_html_patterns:
                experience_html_instruction = experience_html_patterns[style_id]

            # Add specific HTML pattern for education
            if section_key == 'education' and style_id in education_html_patterns:
                education_html_instruction = education_html_patterns[style_id]

            # Add specific HTML pattern for awards
            if section_key == 'awards' and style_id in awards_html_patterns:
                awards_html_instruction = awards_html_patterns[style_id]

    # Add strong achievement warning if that style is selected
    achievement_warning = ""
    if experience_style == 'achievements':
        achievement_warning = """
***** MANDATORY WORK EXPERIENCE FORMAT - ACHIEVEMENT BULLETS *****
DO NOT USE PARAGRAPHS FOR WORK EXPERIENCE. THIS IS NON-NEGOTIABLE.

For EVERY job entry, you MUST use this EXACT format:
<ul style="margin: 8px 0; padding-left: 18px;">
  <li style="margin-bottom: 4px; font-size: 11px;">• Action verb + specific achievement + metric (e.g., "Led team of 15, reducing costs by 30%")</li>
  <li style="margin-bottom: 4px; font-size: 11px;">• Action verb + specific achievement + metric</li>
  <li style="margin-bottom: 4px; font-size: 11px;">• Action verb + specific achievement + metric</li>
</ul>

FORBIDDEN: Long paragraph descriptions. ONLY bullet points allowed.
REQUIRED: 3-5 bullets per job, each starting with action verb, each with metrics.
*****************************************************************
"""

    if section_styles_instruction:
        section_styles_instruction = f"""
CRITICAL - SECTION DESIGN STYLES (MUST FOLLOW EXACTLY):
{section_styles_instruction}

{achievement_warning}

{skills_html_instruction}

{summary_html_instruction}

{experience_html_instruction}

{education_html_instruction}

{awards_html_instruction}
"""
        print(f"[DEBUG] Skills style instruction: {skills_html_instruction[:100] if skills_html_instruction else 'NONE'}...")
        print(f"[DEBUG] Summary style instruction: {summary_html_instruction[:100] if summary_html_instruction else 'NONE'}...")
        print(f"[DEBUG] Experience style instruction: {experience_html_instruction[:100] if experience_html_instruction else 'NONE'}...")
        print(f"[DEBUG] Education style instruction: {education_html_instruction[:100] if education_html_instruction else 'NONE'}...")
        print(f"[DEBUG] Awards style instruction: {awards_html_instruction[:100] if awards_html_instruction else 'NONE'}...")

    # CRITICAL: Global 2-column constraint for all templates
    two_column_constraint = """
***** ABSOLUTE CRITICAL LAYOUT RULE - STRICTLY ENFORCED *****
THE RESUME MUST HAVE EXACTLY 2 COLUMNS - NO MORE, NO LESS:

REQUIRED HTML STRUCTURE:
<div style="display: flex;">
  <div style="width: 35%;">LEFT SIDEBAR</div>
  <div style="width: 65%;">MAIN CONTENT</div>
</div>

SIDEBAR (LEFT, 35%): Contains Skills, Education, Contact info
MAIN CONTENT (RIGHT, 65%): Contains Summary, Work Experience, Awards, Achievements - ALL STACKED VERTICALLY

FORBIDDEN:
- DO NOT create 3, 4, or 5 columns
- DO NOT place sections side-by-side horizontally (except the 2-column layout)
- DO NOT put Awards next to Work Experience horizontally
- DO NOT create a grid layout with multiple sections in a row

ALL SECTIONS IN MAIN CONTENT MUST STACK VERTICALLY:
Summary → Work Experience → Awards → Achievements (one below the other)

THIS IS NON-NEGOTIABLE. VIOLATING THIS RULE WILL MAKE THE RESUME UNUSABLE.
"""

    # Build ATS optimization instructions if job description provided
    ats_optimization = ""
    if job_description:
        ats_optimization = f"""
***** CRITICAL ATS OPTIMIZATION - MUST FOLLOW *****
TARGET JOB DESCRIPTION:
{job_description}

ATS KEYWORD INTEGRATION REQUIREMENTS:
1. EXTRACT all key skills, tools, technologies, and domain terms from the job description above
2. IDENTIFY the industry/domain (e.g., Healthcare, Finance, Tech) and use industry-specific terminology
3. REWRITE the Professional Summary to directly address the job requirements using exact keywords from JD
4. ENHANCE work experience bullet points to include relevant keywords where truthfully applicable
5. ADD any matching skills from the JD to the Skills section (only if candidate has related experience)
6. USE THE EXACT PHRASES from the job description (e.g., if JD says "prior authorization (PA)", use that exact term)

KEYWORD MAPPING STRATEGY:
- If JD mentions "healthcare data" and resume has "data analytics" → reframe as "healthcare data analytics"
- If JD mentions "compliance" and resume has "quality standards" → add "compliance and quality standards"
- If JD mentions "cross-functional teams" and resume has "team management" → use "cross-functional team leadership"
- Mirror the job title terminology in the summary (e.g., "Healthcare Data Analyst" if that's the target role)

PROFESSIONAL SUMMARY MUST:
- Start with target role alignment (e.g., "Experienced Data Analyst with expertise in healthcare analytics...")
- Include 3-5 exact keyword phrases from the job description
- Mention relevant tools/technologies that match the JD requirements
- Be tailored specifically to this job, not generic

CRITICAL: Do NOT fabricate experience. Only reframe existing experience using JD terminology where there's genuine overlap.
This ATS optimization is MANDATORY when job description is provided.
"""

    # Build prompt based on template
    if template_id == 'minimal':
        # Single column minimal design
        prompt = f"""Generate a professional MINIMAL single-column HTML resume. Use clean, simple styling.

TEMPLATE STYLE: Minimal - Single column, clean typography, lots of whitespace
COLORS: Primary: {colors['primary']}, Secondary: {colors['secondary']}, Accent: {colors['accent']}

Create a clean HTML resume with:
- Single column layout centered on page (max-width: 700px)
- Clean header with name centered, subtitle below
- Small circular profile photo (80px) centered above name if photo URL provided
- Thin horizontal lines as section dividers
- Simple, readable typography (font-family: 'Georgia', serif for headings, system-ui for body)
- Generous whitespace and margins
- Minimal use of colors - mostly black/gray text
- Include AWARDS & ACHIEVEMENTS section if data exists

DATA:
Name: {full_name}
Photo URL: {profile_picture_url or 'No photo'}
Email: {profile.email or user.email}
Phone: {profile.phone or 'Not provided'}
Location: {profile.location or 'Not provided'}
Summary: {profile.bio or 'Experienced professional seeking new opportunities.'}

WORK EXPERIENCE:
{career_text or 'No work experience listed'}

EDUCATION:
{education_text or 'No education listed'}

SKILLS: {skills_text}

AWARDS:
{awards_text or 'None listed'}

ACHIEVEMENTS:
{achievements_text or 'None listed'}

{ats_optimization}

{section_order_instruction}
{section_styles_instruction}

CRITICAL - RESUME MUST FIT ON 1-2 PAGES MAX:
1. Use compact styling: body font-size 10-11px, minimal margins/padding
2. Keep sections tight with minimal spacing
3. Include this REQUIRED print CSS:
@media print {{
  * {{ -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }}
  @page {{ margin: 0.3in; size: A4; }}
  body {{ font-size: 10px !important; line-height: 1.3; margin: 0; }}
  h1 {{ font-size: 18px !important; margin: 0 0 5px 0; }}
  h2, h3 {{ font-size: 12px !important; margin: 5px 0 3px 0; }}
  p, li {{ font-size: 10px !important; margin: 2px 0; }}
  section {{ page-break-inside: avoid; }}
}}

Return ONLY complete HTML starting with <!DOCTYPE html>. No markdown."""

    elif template_id == 'executive':
        # Dark elegant executive design
        prompt = f"""Generate an EXECUTIVE style HTML resume. Sophisticated dark theme for senior professionals.

{two_column_constraint}

TEMPLATE STYLE: Executive - Dark elegant theme with gold accents
COLORS: Background: {colors['bg']}, Primary: {colors['primary']}, Accent: {colors['accent']}

Create an elegant HTML resume with:
- Dark background (#111827) with white/gold text
- Centered gold-colored name in header
- Circular profile photo (100px) with gold border in sidebar if photo URL provided
- STRICT Two-column layout: sidebar (30%) | main content (70%)
- Dark sidebar with lighter main content area
- Gold accent lines and borders
- Elegant serif font for headings (Playfair Display or Georgia)
- Awards & Achievements go INSIDE main content column (stacked vertically)
- Professional accomplishment-focused language

DATA:
Name: {full_name}
Photo URL: {profile_picture_url or 'No photo'}
Email: {profile.email or user.email}
Phone: {profile.phone or 'Not provided'}
Location: {profile.location or 'Not provided'}
Summary: {profile.bio or 'Experienced professional seeking new opportunities.'}

WORK EXPERIENCE:
{career_text or 'No work experience listed'}

EDUCATION:
{education_text or 'No education listed'}

SKILLS: {skills_text}

AWARDS:
{awards_text or 'None listed'}

ACHIEVEMENTS:
{achievements_text or 'None listed'}

{ats_optimization}

{section_order_instruction}
{section_styles_instruction}

CRITICAL - RESUME MUST FIT ON 1-2 PAGES MAX:
1. Use compact styling: body font-size 10-11px, minimal margins/padding
2. Keep sections tight with minimal spacing
3. Include this REQUIRED print CSS:
@media print {{
  * {{ -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }}
  @page {{ margin: 0.3in; size: A4; }}
  body {{ font-size: 10px !important; line-height: 1.3; margin: 0; }}
  h1 {{ font-size: 18px !important; margin: 0 0 5px 0; }}
  h2, h3 {{ font-size: 12px !important; margin: 5px 0 3px 0; }}
  p, li {{ font-size: 10px !important; margin: 2px 0; }}
  section {{ page-break-inside: avoid; }}
}}

Return ONLY complete HTML starting with <!DOCTYPE html>. No markdown."""

    elif template_id == 'creative':
        # Colorful creative design
        prompt = f"""Generate a CREATIVE style HTML resume. Colorful sidebar with modern typography.

{two_column_constraint}

TEMPLATE STYLE: Creative - Vibrant teal sidebar, modern layout
COLORS: Primary: {colors['primary']}, Secondary: {colors['secondary']}, Accent: {colors['accent']}, BG: {colors['bg']}

Create a vibrant HTML resume with:
- STRICT Two-column layout: colorful teal/cyan sidebar (35%) | main content (65%)
- White/light main content area
- Circular profile photo (120px) with white border at top of sidebar - USE ACTUAL PHOTO URL if provided
- Skills shown as colorful tags/pills in sidebar
- Modern sans-serif typography (Inter or system-ui)
- Rounded corners on sections
- Awards & Achievements go INSIDE main content column (stacked vertically below Experience)
- Fun but professional appearance

DATA:
Name: {full_name}
Photo URL: {profile_picture_url or 'No photo'}
Email: {profile.email or user.email}
Phone: {profile.phone or 'Not provided'}
Location: {profile.location or 'Not provided'}
Summary: {profile.bio or 'Experienced professional seeking new opportunities.'}

WORK EXPERIENCE:
{career_text or 'No work experience listed'}

EDUCATION:
{education_text or 'No education listed'}

SKILLS: {skills_text}

AWARDS:
{awards_text or 'None listed'}

ACHIEVEMENTS:
{achievements_text or 'None listed'}

{ats_optimization}

{section_order_instruction}
{section_styles_instruction}

CRITICAL - RESUME MUST FIT ON 1-2 PAGES MAX:
1. Use compact styling: body font-size 10-11px, minimal margins/padding
2. Keep sections tight with minimal spacing
3. Include this REQUIRED print CSS:
@media print {{
  * {{ -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }}
  @page {{ margin: 0.3in; size: A4; }}
  body {{ font-size: 10px !important; line-height: 1.3; margin: 0; }}
  h1 {{ font-size: 18px !important; margin: 0 0 5px 0; }}
  h2, h3 {{ font-size: 12px !important; margin: 5px 0 3px 0; }}
  p, li {{ font-size: 10px !important; margin: 2px 0; }}
  section {{ page-break-inside: avoid; }}
}}

Return ONLY complete HTML starting with <!DOCTYPE html>. No markdown."""

    elif template_id == 'tech':
        # Developer/tech focused design
        prompt = f"""Generate a TECH/DEVELOPER style HTML resume. Code-inspired dark theme.

{two_column_constraint}

TEMPLATE STYLE: Tech - Dark theme like a code editor, monospace elements
COLORS: Background: {colors['bg']}, Primary: {colors['primary']}, Accent: {colors['accent']}

Create a developer-focused HTML resume with:
- Dark slate background (#0f172a)
- Terminal/code editor style header with mock window buttons (red/yellow/green dots)
- STRICT Two-column layout: sidebar (30-35%) | main content (65-70%)
- Small circular profile photo (80px) with green border in sidebar if photo URL provided
- Monospace font for headings (JetBrains Mono or Consolas)
- Green accent color for highlights (like terminal green)
- Skills displayed as "code" style tags in sidebar
- Section titles with // comment style or code-like formatting
- Awards & Achievements go INSIDE main content column (stacked vertically)
- Technical, clean appearance

DATA:
Name: {full_name}
Photo URL: {profile_picture_url or 'No photo'}
Email: {profile.email or user.email}
Phone: {profile.phone or 'Not provided'}
Location: {profile.location or 'Not provided'}
Summary: {profile.bio or 'Experienced professional seeking new opportunities.'}

WORK EXPERIENCE:
{career_text or 'No work experience listed'}

EDUCATION:
{education_text or 'No education listed'}

SKILLS: {skills_text}

AWARDS:
{awards_text or 'None listed'}

ACHIEVEMENTS:
{achievements_text or 'None listed'}

{ats_optimization}

{section_order_instruction}
{section_styles_instruction}

CRITICAL - RESUME MUST FIT ON 1-2 PAGES MAX:
1. Use compact styling: body font-size 10-11px, minimal margins/padding
2. Keep sections tight with minimal spacing
3. Include this REQUIRED print CSS:
@media print {{
  * {{ -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }}
  @page {{ margin: 0.3in; size: A4; }}
  body {{ font-size: 10px !important; line-height: 1.3; margin: 0; }}
  h1 {{ font-size: 18px !important; margin: 0 0 5px 0; }}
  h2, h3 {{ font-size: 12px !important; margin: 5px 0 3px 0; }}
  p, li {{ font-size: 10px !important; margin: 2px 0; }}
  section {{ page-break-inside: avoid; }}
}}

Return ONLY complete HTML starting with <!DOCTYPE html>. No markdown."""

    elif template_id == 'modern':
        # Modern purple design
        prompt = f"""Generate a MODERN style HTML resume. Bold purple header with clean sections.

{two_column_constraint}

TEMPLATE STYLE: Modern - Bold purple/indigo theme, contemporary layout
COLORS: Primary: {colors['primary']}, Secondary: {colors['secondary']}, Accent: {colors['accent']}, BG: {colors['bg']}

Create a modern HTML resume with:
- Bold purple/indigo gradient header
- Circular profile photo (100px) with white border in header area if photo URL provided
- Name prominently displayed in white
- Contact info as horizontal pills below name
- Clean white main content area
- STRICT Two-column layout: skills sidebar (30%) | main content (70%)
- Purple accent lines for section dividers
- Skills as purple-tinted pills
- Awards & Achievements go INSIDE main content column (stacked vertically below Experience)
- Contemporary, fresh appearance

DATA:
Name: {full_name}
Photo URL: {profile_picture_url or 'No photo'}
Email: {profile.email or user.email}
Phone: {profile.phone or 'Not provided'}
Location: {profile.location or 'Not provided'}
Summary: {profile.bio or 'Experienced professional seeking new opportunities.'}

WORK EXPERIENCE:
{career_text or 'No work experience listed'}

EDUCATION:
{education_text or 'No education listed'}

SKILLS: {skills_text}

AWARDS:
{awards_text or 'None listed'}

ACHIEVEMENTS:
{achievements_text or 'None listed'}

{ats_optimization}

{section_order_instruction}
{section_styles_instruction}

CRITICAL - RESUME MUST FIT ON 1-2 PAGES MAX:
1. Use compact styling: body font-size 10-11px, minimal margins/padding
2. Keep sections tight with minimal spacing
3. Include this REQUIRED print CSS:
@media print {{
  * {{ -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }}
  @page {{ margin: 0.3in; size: A4; }}
  body {{ font-size: 10px !important; line-height: 1.3; margin: 0; }}
  h1 {{ font-size: 18px !important; margin: 0 0 5px 0; }}
  h2, h3 {{ font-size: 12px !important; margin: 5px 0 3px 0; }}
  p, li {{ font-size: 10px !important; margin: 2px 0; }}
  section {{ page-break-inside: avoid; }}
}}

Return ONLY complete HTML starting with <!DOCTYPE html>. No markdown."""

    elif template_id == 'classic':
        # Classic template with traditional bordered design
        prompt = f"""Generate a CLASSIC style HTML resume. Traditional bordered design with elegant typography.

{two_column_constraint}

TEMPLATE STYLE: Classic - Traditional bordered resume with elegant serif typography
COLORS: Primary: {colors['primary']}, Secondary: {colors['secondary']}, Accent: {colors['accent']}

CRITICAL DESIGN REQUIREMENTS FOR CLASSIC STYLE:
1. OUTER BORDER: The entire resume must have a double-line border frame (border: 3px double #1e3a5f)
2. HEADER: Clean centered header with name in elegant serif font, NO background color, bordered bottom
3. STRICT TWO COLUMNS: Left sidebar (35%) with light blue accent background (#f0f7ff), main content (65%)
4. Awards/Achievements go INSIDE main content column (stacked vertically below Experience)
5. SECTION HEADERS: Elegant with underline decoration, using serif font
6. CIRCULAR PHOTO: If provided, show centered circular photo with border
7. TYPOGRAPHY: Use Georgia or Times New Roman for headings, clean sans-serif for body

CREATE THIS HTML STRUCTURE:
<div class="resume" style="border: 3px double #1e3a5f; padding: 20px; max-width: 210mm; margin: 0 auto; font-family: 'Segoe UI', sans-serif;">
  <div class="header" style="text-align: center; border-bottom: 2px solid #1e3a5f; padding-bottom: 15px; margin-bottom: 20px;">
    <!-- Photo, Name, Contact -->
  </div>
  <div class="content" style="display: flex; gap: 20px;">
    <div class="sidebar" style="width: 35%; background: #f0f7ff; padding: 15px; border-radius: 8px;">
      <!-- Skills, Education -->
    </div>
    <div class="main" style="width: 65%;">
      <!-- Summary, Experience, Awards -->
    </div>
  </div>
</div>

DATA:
Name: {full_name}
Photo URL: {profile_picture_url or 'No photo'}
Email: {profile.email or user.email}
Phone: {profile.phone or 'Not provided'}
Location: {profile.location or 'Not provided'}
Summary: {profile.bio or 'Experienced professional seeking new opportunities.'}

WORK EXPERIENCE:
{career_text or 'No work experience listed'}

EDUCATION:
{education_text or 'No education listed'}

SKILLS: {skills_text}

AWARDS:
{awards_text or 'None listed'}

ACHIEVEMENTS:
{achievements_text or 'None listed'}

{ats_optimization}

{section_order_instruction}
{section_styles_instruction}

Include print CSS for proper A4 printing.
Return ONLY complete HTML starting with <!DOCTYPE html>. No markdown."""

    else:
        # Default professional template - A4 sized, multi-page support
        prompt = f"""Generate a professional HTML resume with proper A4 formatting.

{two_column_constraint}

A4 PAGE REQUIREMENTS:
- Width: 210mm (A4 width)
- Allow content to flow to multiple pages naturally
- Font sizes: Name 22px, Section headers 14px, Body text 11px
- Margins: 15mm padding
- Line height: 1.4

TEMPLATE STYLE: Professional - STRICT Two-column layout
COLORS: Primary: {colors['primary']}, Secondary: {colors['secondary']}, Accent: {colors['accent']}

STRUCTURE:
1. HEADER: Name prominently displayed, contact info below (email | phone | location)
2. Photo: If URL provided, show 80px CIRCULAR photo (border-radius: 50%) in header area - MUST be circular, not square
3. STRICT TWO COLUMNS ONLY: Left sidebar (35%) for Skills & Education. Main area (65%) for Summary, Experience, Awards (ALL stacked vertically)
4. Awards and Achievements must be INSIDE main content column, stacked below Experience - NOT as a separate third column
5. Include ALL work experience with descriptions
6. Page breaks should happen naturally between sections

REQUIRED CSS:
```css
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
@page {{ size: A4; margin: 10mm; }}
html, body {{ width: 210mm; font-family: 'Segoe UI', Arial, sans-serif; font-size: 11px; line-height: 1.4; background: white; }}
.resume {{ width: 210mm; padding: 15mm; background: white; }}
.header {{ text-align: center; padding-bottom: 15px; border-bottom: 3px solid {colors['primary']}; margin-bottom: 20px; }}
.header h1 {{ font-size: 28px; color: {colors['primary']}; margin-bottom: 8px; }}
.contact {{ font-size: 11px; color: #555; }}
.contact span {{ margin: 0 10px; }}
.photo {{ width: 80px; height: 80px; border-radius: 50%; object-fit: cover; border: 3px solid {colors['primary']}; margin-bottom: 10px; }}
.content {{ display: flex; gap: 20px; }}
.sidebar {{ width: 35%; }}
.main {{ width: 65%; }}
.section {{ margin-bottom: 20px; page-break-inside: avoid; }}
.section-title {{ font-size: 14px; font-weight: bold; color: {colors['primary']}; border-bottom: 2px solid {colors['secondary']}; padding-bottom: 5px; margin-bottom: 10px; text-transform: uppercase; }}
.job {{ margin-bottom: 15px; }}
.job-header {{ display: flex; justify-content: space-between; align-items: baseline; }}
.job-title {{ font-size: 13px; font-weight: bold; color: #333; }}
.job-company {{ font-size: 12px; color: {colors['secondary']}; font-weight: 500; }}
.job-date {{ font-size: 10px; color: #888; }}
.job-desc {{ font-size: 11px; margin-top: 5px; color: #444; }}
.job-desc ul {{ margin-left: 15px; margin-top: 5px; }}
.job-desc li {{ margin-bottom: 3px; }}
.skill-tag {{ display: inline-block; background: {colors['secondary']}; color: white; padding: 4px 10px; border-radius: 4px; font-size: 10px; margin: 3px; }}
.edu-item {{ margin-bottom: 10px; }}
.edu-degree {{ font-weight: bold; font-size: 12px; }}
.edu-school {{ color: {colors['secondary']}; font-size: 11px; }}
.edu-year {{ font-size: 10px; color: #888; }}
@media print {{
  * {{ -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }}
  .resume {{ padding: 10mm; min-height: 0 !important; }}
  .section {{ page-break-inside: avoid; }}
  html, body {{ min-height: 0 !important; }}
}}
/* Ensure circular photo */
img {{ border-radius: 50%; object-fit: cover; }}
```

DATA:
Name: {full_name}
Photo URL: {profile_picture_url if profile_picture_url else 'NONE'}
Email: {profile.email or user.email}
Phone: {profile.phone or 'N/A'}
Location: {profile.location or 'N/A'}
{f'Summary: {profile.bio or "Experienced professional."}' if include_summary else 'Summary: DO NOT INCLUDE - user chose to hide this section'}

{f'WORK EXPERIENCE (include ALL):{chr(10)}{career_text or "No experience"}' if include_experience else 'WORK EXPERIENCE: DO NOT INCLUDE - user chose to hide this section'}

{f'EDUCATION:{chr(10)}{education_text or "No education"}' if include_education else 'EDUCATION: DO NOT INCLUDE - user chose to hide this section'}

{f'SKILLS: {skills_text}' if include_skills else 'SKILLS: DO NOT INCLUDE - user chose to hide this section'}

{f'AWARDS: {awards_text or "None"}' if include_awards else 'AWARDS: DO NOT INCLUDE - user chose to hide this section'}

{f'ACHIEVEMENTS: {achievements_text or "None"}' if include_achievements else 'ACHIEVEMENTS: DO NOT INCLUDE - user chose to hide this section'}

IMPORTANT: Only include sections that have data above. If a section says "DO NOT INCLUDE", skip it entirely in the resume.

Return ONLY the complete HTML. No markdown. Include all selected content - let it flow to multiple pages if needed."""

    try:
        message = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}]
        )
        resume_html = message.content[0].text

        # Clean up the response
        resume_html = resume_html.strip()
        if resume_html.startswith("```"):
            resume_html = re.sub(r'^```html?\n?', '', resume_html)
            resume_html = re.sub(r'\n?```$', '', resume_html)

        # Post-process skills section to apply correct style
        skills_style = section_styles.get('skills', 'chips')
        if skills_style != 'chips' and skills:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(resume_html, 'html.parser')

            # Find skills section by looking for "SKILLS" heading
            skills_section = None
            for heading in soup.find_all(['h2', 'h3', 'div']):
                if heading.get_text(strip=True).upper() == 'SKILLS':
                    skills_section = heading.parent if heading.name != 'div' else heading
                    break

            if skills_section:
                # Build skills list with proficiency
                skills_data = []
                for s in skills:
                    prof_pct = s.proficiency * 20 if s.proficiency <= 5 else min(100, s.proficiency)
                    skills_data.append({'name': s.name, 'proficiency': prof_pct})

                # Generate new skills HTML based on style
                if skills_style == 'bars':
                    skills_html = f'<div class="section-title" style="font-size: 14px; font-weight: bold; color: {colors["primary"]}; border-bottom: 2px solid {colors["secondary"]}; padding-bottom: 5px; margin-bottom: 10px; text-transform: uppercase;">SKILLS</div>'
                    for skill in skills_data:
                        skills_html += f'''<div style="margin-bottom: 8px;">
                            <div style="display: flex; justify-content: space-between; font-size: 11px; margin-bottom: 3px;">
                                <span style="font-weight: 500;">{skill['name']}</span>
                                <span style="color: #6b7280;">{skill['proficiency']}%</span>
                            </div>
                            <div style="height: 6px; background: #e5e7eb; border-radius: 3px; overflow: hidden;">
                                <div style="width: {skill['proficiency']}%; height: 100%; background: {colors["primary"]}; border-radius: 3px;"></div>
                            </div>
                        </div>'''
                elif skills_style == 'grid':
                    skills_html = f'<div class="section-title" style="font-size: 14px; font-weight: bold; color: {colors["primary"]}; border-bottom: 2px solid {colors["secondary"]}; padding-bottom: 5px; margin-bottom: 10px; text-transform: uppercase;">SKILLS</div>'
                    skills_html += '<div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 6px;">'
                    for skill in skills_data:
                        skills_html += f'<div style="background: #f3f4f6; padding: 6px 8px; border-radius: 4px; text-align: center; font-size: 10px;">{skill["name"]}</div>'
                    skills_html += '</div>'
                elif skills_style == 'list':
                    skills_html = f'<div class="section-title" style="font-size: 14px; font-weight: bold; color: {colors["primary"]}; border-bottom: 2px solid {colors["secondary"]}; padding-bottom: 5px; margin-bottom: 10px; text-transform: uppercase;">SKILLS</div>'
                    skills_html += '<ul style="list-style: disc; padding-left: 20px; margin: 0;">'
                    for skill in skills_data:
                        skills_html += f'<li style="font-size: 11px; margin-bottom: 2px;">{skill["name"]}</li>'
                    skills_html += '</ul>'
                else:
                    skills_html = None

                if skills_html:
                    # Replace the skills section content
                    new_section = BeautifulSoup(f'<div class="section">{skills_html}</div>', 'html.parser')
                    skills_section.replace_with(new_section)
                    resume_html = str(soup)

        return Response({'resume_html': resume_html})
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def resume_approve(request):
    """Save approved resume as HTML and generate Word document"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup

    profile = get_object_or_404(UserProfile, user=request.user)
    resume_html = request.data.get('resume_html')
    template = request.data.get('template', 'professional')

    if not resume_html:
        return Response({'error': 'No resume HTML provided'}, status=status.HTTP_400_BAD_REQUEST)

    # Save HTML version and template
    profile.ai_resume_html = resume_html
    profile.resume_template = template

    # Convert HTML to Word document
    try:
        soup = BeautifulSoup(resume_html, 'html.parser')
        doc = Document()

        # Get the body content
        body = soup.find('body') or soup

        def process_element(element):
            """Recursively process HTML elements"""
            if element.name is None:
                return

            # Skip script/style tags
            if element.name in ['script', 'style', 'head', 'meta', 'link']:
                return

            # Get direct text (not from children)
            text = element.get_text(strip=True)

            if element.name == 'h1':
                if text:
                    p = doc.add_heading(text, level=0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif element.name == 'h2':
                if text:
                    doc.add_heading(text, level=1)
            elif element.name == 'h3':
                if text:
                    doc.add_heading(text, level=2)
            elif element.name == 'h4':
                if text:
                    doc.add_heading(text, level=3)
            elif element.name == 'li':
                if text:
                    doc.add_paragraph(text, style='List Bullet')
            elif element.name == 'p':
                if text:
                    doc.add_paragraph(text)
            elif element.name in ['div', 'section', 'article', 'header', 'main']:
                # For containers, process children
                for child in element.children:
                    if hasattr(child, 'name'):
                        process_element(child)
            elif element.name == 'ul' or element.name == 'ol':
                # Process list items
                for li in element.find_all('li', recursive=False):
                    li_text = li.get_text(strip=True)
                    if li_text:
                        doc.add_paragraph(li_text, style='List Bullet')
            elif element.name == 'strong' or element.name == 'b':
                # Bold text - add as paragraph
                if text and element.parent.name not in ['p', 'li', 'h1', 'h2', 'h3', 'h4']:
                    doc.add_paragraph(text)

        # Process all direct children of body
        for child in body.children:
            if hasattr(child, 'name') and child.name:
                process_element(child)

        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Save to model
        filename = f"resume_{request.user.username}.docx"
        profile.ai_resume.save(filename, ContentFile(docx_buffer.read()), save=False)
        profile.save()

        return Response({
            'success': True,
            'message': 'Resume saved successfully',
            'resume_url': profile.ai_resume.url
        })
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def resume_download(request):
    """Download the approved resume as Word document"""
    profile = get_object_or_404(UserProfile, user=request.user)

    if not profile.ai_resume or not profile.ai_resume.name:
        return Response({'error': 'No resume available'}, status=status.HTTP_404_NOT_FOUND)

    return FileResponse(
        profile.ai_resume.open('rb'),
        as_attachment=True,
        filename=f"resume_{request.user.username}.docx"
    )


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def resume_download_pdf(request):
    """Download the approved resume as PDF (admin user)"""
    from xhtml2pdf import pisa

    profile = get_object_or_404(UserProfile, user=request.user)

    if not profile.ai_resume_html:
        return Response({'error': 'No resume available'}, status=status.HTTP_404_NOT_FOUND)

    # Get template CSS and inject it
    template = getattr(profile, 'resume_template', 'professional') or 'professional'
    template_css = RESUME_TEMPLATE_CSS.get(template, RESUME_TEMPLATE_CSS['professional'])

    resume_html = profile.ai_resume_html

    # Add PDF-specific styles to prevent blank pages
    pdf_styles = """
    @page {
        size: A4;
        margin: 0.5cm;
    }
    html, body {
        margin: 0;
        padding: 0;
        height: auto !important;
        overflow: visible !important;
    }
    .resume, .resume-container, #resume {
        page-break-after: avoid;
        page-break-inside: avoid;
    }
    """

    if '</style>' in resume_html:
        resume_html = resume_html.replace('</style>', f'{template_css}\n{pdf_styles}\n</style>')
    elif '<head>' in resume_html:
        resume_html = resume_html.replace('<head>', f'<head><style>{pdf_styles}</style>')

    # Embed profile picture as base64 for PDF generation
    resume_html = embed_profile_picture_in_html(resume_html, profile)

    # Convert HTML to PDF
    pdf_buffer = BytesIO()
    pisa_status = pisa.CreatePDF(resume_html, dest=pdf_buffer)

    if pisa_status.err:
        return Response({'error': 'PDF generation failed'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    pdf_buffer.seek(0)
    response = HttpResponse(pdf_buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="resume_{request.user.username}.pdf"'
    return response


@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
def resume_download_docx(request):
    """Download the resume as Word document (.docx) from HTML"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from bs4 import BeautifulSoup

    # Get profile first
    profile = get_object_or_404(UserProfile, user=request.user)

    # Use HTML from POST body if provided, otherwise fallback to database
    html_content = None
    if request.method == 'POST' and request.data.get('html'):
        html_content = request.data.get('html')

    # Fallback to database if POST html is empty
    if not html_content:
        html_content = profile.ai_resume_html

    if not html_content:
        return Response({'error': 'No resume available. Please generate one first.'}, status=status.HTTP_404_NOT_FOUND)

    # Parse the HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Try to extract structured content
    # Get name from h1 or header
    name_elem = soup.find('h1')
    if name_elem:
        heading = doc.add_heading(name_elem.get_text(strip=True), 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Get contact info
    contact_div = soup.find(class_=['contact', 'contact-info', 'header'])
    if contact_div:
        contact_text = contact_div.get_text(separator=' | ', strip=True)
        if contact_text and name_elem:
            contact_text = contact_text.replace(name_elem.get_text(strip=True), '').strip(' |')
        if contact_text:
            p = doc.add_paragraph(contact_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Find all sections
    sections = soup.find_all(['section', 'div'], class_=lambda x: x and ('section' in str(x).lower() if x else False))

    # If no sections found, try to find headings
    if not sections:
        sections = soup.find_all(['h2', 'h3'])

    processed_headings = set()

    for section in sections:
        # Get section heading
        heading_elem = section.find(['h2', 'h3']) if section.name != 'h2' and section.name != 'h3' else section
        if heading_elem:
            heading_text = heading_elem.get_text(strip=True).upper()
            if heading_text in processed_headings or not heading_text:
                continue
            processed_headings.add(heading_text)

            # Add section heading
            doc.add_heading(heading_text, level=1)

            # Get section content
            if section.name in ['h2', 'h3']:
                # Find next sibling content
                next_elem = section.find_next_sibling()
                if next_elem:
                    content = next_elem.get_text(separator='\n', strip=True)
                    if content:
                        doc.add_paragraph(content)
            else:
                # Get all text content except the heading
                for child in section.children:
                    if child == heading_elem or (hasattr(child, 'name') and child.name in ['h2', 'h3']):
                        continue
                    if hasattr(child, 'get_text'):
                        text = child.get_text(separator='\n', strip=True)
                        if text:
                            # Check for bullet points
                            if child.name in ['ul', 'ol']:
                                for li in child.find_all('li'):
                                    li_text = li.get_text(strip=True)
                                    if li_text:
                                        doc.add_paragraph(li_text, style='List Bullet')
                            else:
                                doc.add_paragraph(text)

    # If no sections were processed, try a simpler approach
    if not processed_headings:
        # Just extract all text with basic formatting
        body = soup.find('body') or soup
        for elem in body.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'li', 'div']):
            text = elem.get_text(strip=True)
            if not text:
                continue
            if elem.name == 'h1':
                doc.add_heading(text, 0)
            elif elem.name in ['h2', 'h3']:
                doc.add_heading(text, 1)
            elif elem.name == 'li':
                doc.add_paragraph(text, style='List Bullet')
            elif elem.name in ['p', 'div'] and len(text) > 10:
                doc.add_paragraph(text)

    # Save to buffer
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    response = HttpResponse(
        docx_buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="resume_{request.user.username}.docx"'
    return response


@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def public_resume_download(request, username):
    """Public endpoint to download user's resume as PDF"""
    from xhtml2pdf import pisa

    profile = get_object_or_404(UserProfile, user__username=username)

    # Try AI-generated HTML resume -> convert to PDF
    if profile.ai_resume_html:
        try:
            # Get template CSS and inject it
            template = getattr(profile, 'resume_template', 'professional') or 'professional'
            template_css = RESUME_TEMPLATE_CSS.get(template, RESUME_TEMPLATE_CSS['professional'])

            resume_html = profile.ai_resume_html
            if '</style>' in resume_html:
                resume_html = resume_html.replace('</style>', f'{template_css}\n</style>')

            # Embed profile picture as base64 for PDF generation
            resume_html = embed_profile_picture_in_html(resume_html, profile)

            # Convert HTML to PDF
            pdf_buffer = BytesIO()
            pisa_status = pisa.CreatePDF(resume_html, dest=pdf_buffer)

            if not pisa_status.err:
                pdf_buffer.seek(0)
                response = HttpResponse(pdf_buffer.read(), content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="resume_{username}.pdf"'
                return response
        except Exception as e:
            print(f"PDF conversion error: {e}")
            # Continue to fallbacks
            pass

    # Option 2: Try AI-generated docx file
    if profile.ai_resume and profile.ai_resume.name:
        return FileResponse(
            profile.ai_resume.open('rb'),
            as_attachment=True,
            filename=f"resume_{username}.docx"
        )

    # Option 3: Try user's uploaded CV file
    if profile.cv and profile.cv.name:
        # Get the original file extension
        import os
        ext = os.path.splitext(profile.cv.name)[1] or '.pdf'
        return FileResponse(
            profile.cv.open('rb'),
            as_attachment=True,
            filename=f"cv_{username}{ext}"
        )

    return Response({'error': 'No resume available'}, status=status.HTTP_404_NOT_FOUND)


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def resume_html(request):
    """Get the HTML version of the approved resume"""
    profile = get_object_or_404(UserProfile, user=request.user)

    if not profile.ai_resume_html:
        return Response({'error': 'No resume HTML available'}, status=status.HTTP_404_NOT_FOUND)

    template = getattr(profile, 'resume_template', 'professional') or 'professional'

    # Inject print CSS to preserve colors and enforce A4 sizing
    print_css = """
<style id="print-force-css">
@media print {
  * { -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; color-adjust: exact !important; }
  @page { size: A4; margin: 0; }
  html, body { width: 210mm !important; margin: 0 !important; padding: 0 !important; }
  .resume, .resume-container, main { width: 210mm !important; }
}
</style>
"""

    resume_html_content = profile.ai_resume_html
    # Inject before </head> or before </html>
    if '</head>' in resume_html_content:
        resume_html_content = resume_html_content.replace('</head>', print_css + '</head>')
    elif '</html>' in resume_html_content:
        resume_html_content = resume_html_content.replace('</html>', print_css + '</html>')
    else:
        resume_html_content = resume_html_content + print_css

    # Embed profile picture as base64 so it works when printing/saving as PDF
    resume_html_content = embed_profile_picture_in_html(resume_html_content, profile)

    return Response({
        'resume_html': resume_html_content,
        'template': template
    })


@api_view(['PUT'])
@permission_classes([permissions.IsAuthenticated])
def update_resume_html(request):
    """Update the HTML content of the resume (for inline editing)"""
    profile = get_object_or_404(UserProfile, user=request.user)

    resume_html = request.data.get('resume_html')
    if not resume_html:
        return Response({'error': 'No resume_html provided'}, status=status.HTTP_400_BAD_REQUEST)

    # Save the updated HTML
    profile.ai_resume_html = resume_html
    profile.save(update_fields=['ai_resume_html'])

    return Response({
        'message': 'Resume updated successfully',
        'resume_html': resume_html
    })


@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def public_resume_html(request, username):
    """Public endpoint to get the HTML version of user's resume with template styling"""
    profile = get_object_or_404(UserProfile, user__username__iexact=username)

    if not profile.ai_resume_html:
        return Response({'error': 'No resume available'}, status=status.HTTP_404_NOT_FOUND)

    # Get template CSS and inject it
    template = getattr(profile, 'resume_template', 'professional') or 'professional'
    template_css = RESUME_TEMPLATE_CSS.get(template, RESUME_TEMPLATE_CSS['professional'])

    # Inject template CSS into the HTML
    resume_html_content = profile.ai_resume_html
    if '</style>' in resume_html_content:
        resume_html_content = resume_html_content.replace('</style>', f'{template_css}\n</style>')

    # Inject print CSS to preserve colors and enforce A4 sizing
    print_css = """
<style id="print-force-css">
@media print {
  * { -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; color-adjust: exact !important; }
  @page { size: A4; margin: 0; }
  html, body { width: 210mm !important; margin: 0 !important; padding: 0 !important; }
  .resume, .resume-container, main { width: 210mm !important; }
}
</style>
"""
    # Inject before </head> or before </html>
    if '</head>' in resume_html_content:
        resume_html_content = resume_html_content.replace('</head>', print_css + '</head>')
    elif '</html>' in resume_html_content:
        resume_html_content = resume_html_content.replace('</html>', print_css + '</html>')
    else:
        resume_html_content = resume_html_content + print_css

    # Embed profile picture as base64 so it works when printing/saving as PDF
    resume_html_content = embed_profile_picture_in_html(resume_html_content, profile)

    return Response({
        'resume_html': resume_html_content,
        'template': template
    })


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def check_cv_status(request):
    """Check if user has a CV uploaded in their profile"""
    profile = get_object_or_404(UserProfile, user=request.user)
    has_cv = bool(profile.cv and profile.cv.name)
    cv_name = os.path.basename(profile.cv.name) if has_cv else None
    return Response({
        'has_cv': has_cv,
        'cv_name': cv_name
    })


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def resume_parse_existing(request):
    """Parse the CV already uploaded in user's profile"""
    from PyPDF2 import PdfReader
    from docx import Document
    import json

    profile = get_object_or_404(UserProfile, user=request.user)

    if not profile.cv or not profile.cv.name:
        return Response({'error': 'No CV uploaded in your profile. Please upload a CV in the Profile section first.'},
                       status=status.HTTP_400_BAD_REQUEST)

    filename = profile.cv.name.lower()
    extracted_text = ""

    try:
        # Extract text based on file type
        if filename.endswith('.pdf'):
            reader = PdfReader(profile.cv.open('rb'))
            for page in reader.pages:
                extracted_text += page.extract_text() + "\n"
        elif filename.endswith('.docx'):
            doc = Document(profile.cv.open('rb'))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
        else:
            return Response({'error': 'Unsupported file format. Please upload PDF or DOCX.'},
                          status=status.HTTP_400_BAD_REQUEST)

        if not extracted_text.strip():
            return Response({'error': 'Could not extract text from the CV.'},
                          status=status.HTTP_400_BAD_REQUEST)

        # Use Claude to parse the resume into structured data
        parse_prompt = f"""Parse this resume text and extract structured information.
Return a JSON object with the following structure (use null for missing fields):

{{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, Country",
    "summary": "Professional summary/bio",
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "start_date": "YYYY-MM",
            "end_date": "YYYY-MM or Present",
            "description": "Job description and achievements"
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "institution": "School/University Name",
            "start_date": "YYYY",
            "end_date": "YYYY",
            "description": "Additional details"
        }}
    ],
    "skills": ["skill1", "skill2", "skill3"]
}}

RESUME TEXT:
{extracted_text}

Return ONLY the JSON object, no explanation or markdown."""

        message = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": parse_prompt}]
        )

        response_text = message.content[0].text.strip()

        # Clean up response if it has markdown
        if response_text.startswith("```"):
            response_text = re.sub(r'^```json?\n?', '', response_text)
            response_text = re.sub(r'\n?```$', '', response_text)

        # Parse JSON
        parsed_data = json.loads(response_text)

        return Response({
            'success': True,
            'parsed_data': parsed_data,
            'raw_text': extracted_text[:2000]
        })

    except json.JSONDecodeError as e:
        return Response({'error': f'Failed to parse AI response: {str(e)}'},
                       status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def resume_parse(request):
    """Parse uploaded PDF/Word resume and extract structured data using Claude AI"""
    from PyPDF2 import PdfReader
    from docx import Document
    import json

    uploaded_file = request.FILES.get('resume')
    if not uploaded_file:
        return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

    filename = uploaded_file.name.lower()
    extracted_text = ""

    try:
        # Extract text based on file type
        if filename.endswith('.pdf'):
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                extracted_text += page.extract_text() + "\n"
        elif filename.endswith('.docx'):
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
        else:
            return Response({'error': 'Unsupported file format. Please upload PDF or DOCX.'},
                          status=status.HTTP_400_BAD_REQUEST)

        if not extracted_text.strip():
            return Response({'error': 'Could not extract text from the file.'},
                          status=status.HTTP_400_BAD_REQUEST)

        # Use Claude to parse the resume into structured data
        parse_prompt = f"""Parse this resume text and extract structured information.
Return a JSON object with the following structure (use null for missing fields):

{{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, Country",
    "summary": "Professional summary/bio",
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "start_date": "YYYY-MM",
            "end_date": "YYYY-MM or Present",
            "description": "Job description and achievements"
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "institution": "School/University Name",
            "start_date": "YYYY",
            "end_date": "YYYY",
            "description": "Additional details"
        }}
    ],
    "skills": ["skill1", "skill2", "skill3"]
}}

RESUME TEXT:
{extracted_text}

Return ONLY the JSON object, no explanation or markdown."""

        message = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": parse_prompt}]
        )

        response_text = message.content[0].text.strip()

        # Clean up response if it has markdown
        if response_text.startswith("```"):
            response_text = re.sub(r'^```json?\n?', '', response_text)
            response_text = re.sub(r'\n?```$', '', response_text)

        # Parse JSON
        parsed_data = json.loads(response_text)

        return Response({
            'success': True,
            'parsed_data': parsed_data,
            'raw_text': extracted_text[:2000]  # First 2000 chars for reference
        })

    except json.JSONDecodeError as e:
        return Response({'error': f'Failed to parse AI response: {str(e)}'},
                       status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def resume_import(request):
    """Import parsed resume data into user's profile, career, education, and skills"""
    from career.models import Career
    from education.models import Education
    from home.models import Skill
    import json

    user = request.user
    profile = get_object_or_404(UserProfile, user=user)
    data = request.data.get('parsed_data')

    if not data:
        return Response({'error': 'No data provided'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        # Update profile
        if data.get('email'):
            profile.email = data['email']
        if data.get('phone'):
            profile.phone = data['phone']
        if data.get('location'):
            profile.location = data['location']
        if data.get('summary'):
            profile.bio = data['summary']
        profile.save()

        # Update user name
        if data.get('name'):
            name_parts = data['name'].split(' ', 1)
            user.first_name = name_parts[0]
            user.last_name = name_parts[1] if len(name_parts) > 1 else ''
            user.save()

        # Import experience as Career entries
        imported_careers = 0
        for exp in data.get('experience', []):
            if exp.get('title') and exp.get('company'):
                Career.objects.create(
                    user=user,
                    title=exp['title'],
                    company=exp['company'],
                    start_date=exp.get('start_date', '2020-01'),
                    end_date=exp.get('end_date') if exp.get('end_date') != 'Present' else None,
                    description=exp.get('description', '')
                )
                imported_careers += 1

        # Import education
        imported_education = 0
        for edu in data.get('education', []):
            if edu.get('degree') and edu.get('institution'):
                Education.objects.create(
                    user=user,
                    degree=edu['degree'],
                    institution=edu['institution'],
                    start_date=edu.get('start_date', '2015'),
                    end_date=edu.get('end_date'),
                    description=edu.get('description', '')
                )
                imported_education += 1

        # Import skills
        imported_skills = 0
        for skill_name in data.get('skills', []):
            if skill_name:
                Skill.objects.get_or_create(user=user, name=skill_name)
                imported_skills += 1

        return Response({
            'success': True,
            'message': f'Imported {imported_careers} careers, {imported_education} education entries, and {imported_skills} skills',
            'imported': {
                'careers': imported_careers,
                'education': imported_education,
                'skills': imported_skills
            }
        })

    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ProfileDesignView(APIView):
    """API view for getting and updating profile design configuration"""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        """Get the profile design config for the authenticated user"""
        design, created = ProfileDesign.objects.get_or_create(user=request.user)
        serializer = ProfileDesignSerializer(design)
        return Response(serializer.data)

    def put(self, request):
        """Update the profile design config for the authenticated user"""
        design, created = ProfileDesign.objects.get_or_create(user=request.user)
        serializer = ProfileDesignSerializer(design, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response({'message': 'Profile design updated successfully'})
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def check_onboarding_status(request):
    """Check if user has completed onboarding (CV uploaded or skipped)"""
    profile, created = UserProfile.objects.get_or_create(user=request.user)
    has_cv = bool(profile.cv and profile.cv.name)
    return Response({
        'onboarding_completed': profile.onboarding_completed,
        'has_cv': has_cv,
        'needs_onboarding': not profile.onboarding_completed and not has_cv
    })


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def upload_cv(request):
    """Upload CV, parse with AI, and auto-fill profile data"""
    import logging
    from PyPDF2 import PdfReader
    from docx import Document
    from career.models import CareerEntry, Skill
    from education.models import EducationEntry
    from home.models import HomePageContent
    from achievements.models import Award
    import json

    logger = logging.getLogger('profiles')
    logger.info("[CV Upload] Starting CV upload process...")
    user = request.user
    logger.info(f"[CV Upload] User: {user.username}")
    profile, created = UserProfile.objects.get_or_create(user=user)

    cv_file = request.FILES.get('cv')
    if not cv_file:
        profile.onboarding_completed = True
        profile.save()
        return Response({
            'success': True,
            'message': 'Onboarding completed',
            'has_cv': False
        })

    # Save the CV file
    profile.cv = cv_file
    profile.save()

    # Extract text from CV
    filename = cv_file.name.lower()
    extracted_text = ""

    try:
        cv_file.seek(0)  # Reset file pointer
        if filename.endswith('.pdf'):
            reader = PdfReader(cv_file)
            for page in reader.pages:
                extracted_text += page.extract_text() + "\n"
        elif filename.endswith('.docx'):
            doc = Document(cv_file)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"

        if not extracted_text.strip():
            logger.error("[CV Upload] ERROR: Could not extract any text from CV")
            profile.onboarding_completed = True
            profile.save()
            return Response({
                'success': True,
                'message': 'CV uploaded but could not extract text',
                'has_cv': True
            })

        logger.info(f"[CV Upload] Extracted {len(extracted_text)} characters from CV")
        # Safely print without emoji issues on Windows
        safe_text = extracted_text[:500].encode('ascii', 'replace').decode('ascii')
        logger.info(f"[CV Upload] First 500 chars: {safe_text}")

        # Use Claude to parse the CV
        parse_prompt = f"""Parse this resume/CV and extract ALL information.
Return a JSON object with this structure (use null for missing fields):

{{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, Country",
    "summary": "Professional summary/objective (2-3 sentences)",
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "year": "2020-2023",
            "description": "Job responsibilities and achievements"
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "university": "University Name",
            "year": "2015-2019",
            "description": "Additional details"
        }}
    ],
    "skills": ["skill1", "skill2", "skill3"],
    "awards": [
        {{
            "title": "Award/Honor Name",
            "organization": "Issuing Organization",
            "year": "2022",
            "description": "Details about the award"
        }}
    ]
}}

RESUME TEXT:
{extracted_text}

Return ONLY the JSON object, no explanation."""

        message = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": parse_prompt}]
        )

        response_text = message.content[0].text.strip()
        if response_text.startswith("```"):
            response_text = re.sub(r'^```json?\n?', '', response_text)
            response_text = re.sub(r'\n?```$', '', response_text)

        parsed_data = json.loads(response_text)
        logger.info("[CV Upload] Successfully parsed CV data:")
        logger.info(f"[CV Upload]   Name: {parsed_data.get('name')}")
        logger.info(f"[CV Upload]   Email: {parsed_data.get('email')}")
        logger.info(f"[CV Upload]   Experience: {len(parsed_data.get('experience', []))} entries")
        logger.info(f"[CV Upload]   Education: {len(parsed_data.get('education', []))} entries")
        logger.info(f"[CV Upload]   Skills: {len(parsed_data.get('skills', []))} entries")
        logger.info(f"[CV Upload]   Awards: {len(parsed_data.get('awards', []))} entries")

        # Auto-fill profile data
        if parsed_data.get('email'):
            profile.email = parsed_data['email']
        if parsed_data.get('phone'):
            profile.phone = parsed_data['phone']
        if parsed_data.get('location'):
            profile.location = parsed_data['location']
        if parsed_data.get('summary'):
            profile.bio = parsed_data['summary']
        profile.onboarding_completed = True
        profile.save()

        # Update user name
        if parsed_data.get('name'):
            name_parts = parsed_data['name'].split(' ', 1)
            user.first_name = name_parts[0]
            user.last_name = name_parts[1] if len(name_parts) > 1 else ''
            user.save()

        # Update home page with name and summary
        home, created = HomePageContent.objects.get_or_create(user=user)
        logger.info(f"[CV Upload] HomePageContent created: {created}")
        if parsed_data.get('name'):
            home.title = parsed_data['name']
            logger.info(f"[CV Upload] Set home.title = {parsed_data['name']}")
        if parsed_data.get('summary'):
            home.description = parsed_data['summary']
            # Safe log without emoji issues
            safe_summary = parsed_data['summary'][:100].encode('ascii', 'replace').decode('ascii')
            logger.info(f"[CV Upload] Set home.description = {safe_summary}...")
        home.save()
        logger.info(f"[CV Upload] Home saved! ID: {home.id}, Title: {home.title}")

        # Import career/experience
        imported_careers = 0
        for exp in parsed_data.get('experience', []):
            if exp.get('title') and exp.get('company'):
                CareerEntry.objects.create(
                    user=user,
                    title=exp['title'],
                    company=exp['company'],
                    year=exp.get('year') or '',
                    description=exp.get('description') or ''
                )
                imported_careers += 1

        # Import education
        imported_education = 0
        for edu in parsed_data.get('education', []):
            if edu.get('degree') and edu.get('university'):
                EducationEntry.objects.create(
                    user=user,
                    degree=edu['degree'],
                    university=edu['university'],
                    year=edu.get('year') or '',
                    description=edu.get('description') or ''
                )
                imported_education += 1

        # Import skills
        imported_skills = 0
        for skill_name in parsed_data.get('skills', []):
            if skill_name:
                Skill.objects.get_or_create(user=user, name=skill_name, defaults={'proficiency': 70})
                imported_skills += 1

        # Import awards
        imported_awards = 0
        for award in parsed_data.get('awards', []):
            if award.get('title'):
                Award.objects.create(
                    user=user,
                    title=award['title'],
                    organization=award.get('organization') or '',
                    year=award.get('year') or '',
                    description=award.get('description') or ''
                )
                imported_awards += 1

        return Response({
            'success': True,
            'message': f'CV parsed! Imported: {imported_careers} jobs, {imported_education} education, {imported_skills} skills, {imported_awards} awards',
            'has_cv': True,
            'parsed_data': parsed_data
        })

    except json.JSONDecodeError as e:
        logger.error(f"[CV Upload] ERROR: JSON decode failed: {e}")
        if 'response_text' in dir():
            safe_resp = response_text[:500].encode('ascii', 'replace').decode('ascii')
            logger.error(f"[CV Upload] Response text was: {safe_resp}")
        profile.onboarding_completed = True
        profile.save()
        return Response({
            'success': True,
            'message': 'CV uploaded but parsing failed',
            'has_cv': True,
            'error': str(e)
        })
    except Exception as e:
        import traceback
        logger.error(f"[CV Upload] ERROR: {type(e).__name__}: {e}")
        logger.error(f"[CV Upload] Traceback: {traceback.format_exc()}")
        profile.onboarding_completed = True
        profile.save()
        return Response({
            'success': True,
            'message': 'CV uploaded',
            'has_cv': True,
            'error': str(e)
        })


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def skip_onboarding(request):
    """Skip CV upload and mark onboarding as completed"""
    profile, created = UserProfile.objects.get_or_create(user=request.user)
    profile.onboarding_completed = True
    profile.save()

    return Response({
        'success': True,
        'message': 'Onboarding skipped'
    })


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def parse_cv_and_fill(request):
    """Parse the CV already in profile and auto-fill profile, career, education, skills, awards"""
    import logging
    from PyPDF2 import PdfReader
    from docx import Document
    from career.models import CareerEntry, Skill
    from education.models import EducationEntry
    from home.models import HomePageContent
    from achievements.models import Award
    import json

    import sys
    logger = logging.getLogger('profiles')

    print("[Re-parse CV] Starting re-parse process...", file=sys.stderr)
    print(f"[Re-parse CV] API Key present: {bool(CLAUDE_API_KEY)}", file=sys.stderr)
    print(f"[Re-parse CV] API Key starts with: {CLAUDE_API_KEY[:20] if CLAUDE_API_KEY else 'EMPTY'}...", file=sys.stderr)

    user = request.user
    print(f"[Re-parse CV] User: {user.username}", file=sys.stderr)
    profile = get_object_or_404(UserProfile, user=user)

    if not profile.cv or not profile.cv.name:
        logger.error("[Re-parse CV] No CV found in profile")
        return Response({
            'success': False,
            'message': 'No CV uploaded. Please upload a CV first.'
        }, status=status.HTTP_400_BAD_REQUEST)

    logger.info(f"[Re-parse CV] CV file: {profile.cv.name}")
    filename = profile.cv.name.lower()
    extracted_text = ""

    try:
        # Extract text
        if filename.endswith('.pdf'):
            logger.info("[Re-parse CV] Reading PDF file...")
            reader = PdfReader(profile.cv.open('rb'))
            for page in reader.pages:
                extracted_text += page.extract_text() + "\n"
        elif filename.endswith('.docx'):
            doc = Document(profile.cv.open('rb'))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
        else:
            return Response({
                'success': False,
                'message': 'Unsupported file format. Please upload PDF or DOCX.'
            }, status=status.HTTP_400_BAD_REQUEST)

        if not extracted_text.strip():
            return Response({
                'success': False,
                'message': 'Could not extract text from CV'
            })

        # Parse with Claude
        parse_prompt = f"""Parse this resume/CV and extract ALL information.
Return a JSON object with this structure (use null for missing fields):

{{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, Country",
    "summary": "Professional summary/objective (2-3 sentences)",
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "year": "2020-2023",
            "description": "Job responsibilities and achievements"
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "university": "University Name",
            "year": "2015-2019",
            "description": "Additional details"
        }}
    ],
    "skills": ["skill1", "skill2", "skill3"],
    "awards": [
        {{
            "title": "Award/Honor Name",
            "organization": "Issuing Organization",
            "year": "2022",
            "description": "Details about the award"
        }}
    ]
}}

RESUME TEXT:
{extracted_text}

Return ONLY the JSON object, no explanation."""

        message = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": parse_prompt}]
        )

        response_text = message.content[0].text.strip()
        if response_text.startswith("```"):
            response_text = re.sub(r'^```json?\n?', '', response_text)
            response_text = re.sub(r'\n?```$', '', response_text)

        parsed_data = json.loads(response_text)

        # Auto-fill profile
        if parsed_data.get('email'):
            profile.email = parsed_data['email']
        if parsed_data.get('phone'):
            profile.phone = parsed_data['phone']
        if parsed_data.get('location'):
            profile.location = parsed_data['location']
        if parsed_data.get('summary'):
            profile.bio = parsed_data['summary']
        profile.save()

        # Update user name
        if parsed_data.get('name'):
            name_parts = parsed_data['name'].split(' ', 1)
            user.first_name = name_parts[0]
            user.last_name = name_parts[1] if len(name_parts) > 1 else ''
            user.save()

        # Update home page
        home, _ = HomePageContent.objects.get_or_create(user=user)
        if parsed_data.get('name'):
            home.title = parsed_data['name']
        if parsed_data.get('summary'):
            home.description = parsed_data['summary']
        home.save()

        # Import career/experience
        imported_careers = 0
        for exp in parsed_data.get('experience', []):
            if exp.get('title') and exp.get('company'):
                CareerEntry.objects.create(
                    user=user,
                    title=exp['title'],
                    company=exp['company'],
                    year=exp.get('year') or '',
                    description=exp.get('description') or ''
                )
                imported_careers += 1

        # Import education
        imported_education = 0
        for edu in parsed_data.get('education', []):
            if edu.get('degree') and edu.get('university'):
                EducationEntry.objects.create(
                    user=user,
                    degree=edu['degree'],
                    university=edu['university'],
                    year=edu.get('year') or '',
                    description=edu.get('description') or ''
                )
                imported_education += 1

        # Import skills
        imported_skills = 0
        for skill_name in parsed_data.get('skills', []):
            if skill_name:
                Skill.objects.get_or_create(user=user, name=skill_name, defaults={'proficiency': 70})
                imported_skills += 1

        # Import awards
        imported_awards = 0
        for award in parsed_data.get('awards', []):
            if award.get('title'):
                Award.objects.create(
                    user=user,
                    title=award['title'],
                    organization=award.get('organization') or '',
                    year=award.get('year') or '',
                    description=award.get('description') or ''
                )
                imported_awards += 1

        return Response({
            'success': True,
            'message': f'Imported: {imported_careers} jobs, {imported_education} education, {imported_skills} skills, {imported_awards} awards',
            'parsed_data': parsed_data
        })

    except json.JSONDecodeError as e:
        print(f"[Re-parse CV] JSON decode error: {e}", file=sys.stderr)
        return Response({
            'success': False,
            'message': 'CV parsed but failed to extract structured data',
            'error': str(e)
        })
    except Exception as e:
        import traceback
        print(f"[Re-parse CV] ERROR: {type(e).__name__}: {e}", file=sys.stderr)
        print(f"[Re-parse CV] Traceback: {traceback.format_exc()}", file=sys.stderr)
        return Response({
            'success': False,
            'message': 'Error parsing CV',
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ========== Profile Notes API Views ==========

class ProfileNoteViewSet(viewsets.ModelViewSet):
    """CRUD operations for Profile Notes (admin feature)"""
    serializer_class = ProfileNoteSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return ProfileNote.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def public_profile_notes(request, username):
    """Get visible notes for a public profile"""
    from django.contrib.auth import get_user_model
    User = get_user_model()
    user = get_object_or_404(User, username=username)
    notes = ProfileNote.objects.filter(user=user, is_visible=True)
    serializer = ProfileNoteSerializer(notes, many=True)
    return Response(serializer.data)
