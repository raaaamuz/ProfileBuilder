# views.py

from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
import google.generativeai as genai
import logging

# Configure Google Gemini API
GOOGLE_API_KEY = "AIzaSyCNQ9vdZnYyKsyRW8SxpeS6alLNFQ372mw"
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash-lite')

# Configure logger
logger = logging.getLogger(__name__)

@api_view(["POST"])
@permission_classes([IsAuthenticated])
def improve_description(request):
    """
    Receives a description and returns an improved version using Gemini.
    """
    # Get the description from the request payload
    description = request.data.get("description", "")
    print("Received description:", description)

    if not description:
        return Response({"error": "No description provided"}, status=400)

    prompt = (
        "You are a professional content editor. Rephrase the following description with minimal changes. "
        "Your output should be almost identical in length and detail to the original text, only fixing grammar and clarity. "
        "Do not add any extra information or generic details. Provide exactly 3 distinct variations that retain the original meaning. "
        "Format your response as three complete paragraphs without numbering or labels.\n\n"
        f"Original Description:\n{description}\n\n"
        "Rephrase the above text in 5 variations:, if you received bulletin points, provide variation for each bulletin point"
    )

    try:
        response = model.generate_content(prompt)
        improved_text = response.text.strip()
        variations = improved_text.split("\n\n")[:3]  # Ensure we extract exactly 3 variations

        return Response({"variations": variations})

    except Exception as e:
        return Response({"error": str(e)}, status=500)


@api_view(["POST"])
@permission_classes([IsAuthenticated])
def generate_keywords(request):
    """
    Receives content and returns a list of relevant keywords for searching images on Pexels.
    """
    content = request.data.get("content", "")
    print("Received content for keyword generation:", content)

    if not content:
        return Response({"error": "No content provided"}, status=400)

    prompt = (
        "You are an expert in generating image search keywords. Given the following content, "
        "create a list of 10 relevant keywords that capture the essence, mood, and important themes of the text. "
        "Return the keywords as a comma-separated list without any extra explanation or text.\n\n"
        f"Content:\n{content}"
    )

    try:
        response = model.generate_content(prompt)
        keywords_text = response.text.strip()
        # Split keywords by comma and remove any extra whitespace
        keywords = [keyword.strip() for keyword in keywords_text.split(",") if keyword.strip()]
        return Response({"keywords": keywords})

    except Exception as e:
        return Response({"error": str(e)}, status=500)


import json
import re
import random

# Helper functions for color contrast
def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def get_luminance(hex_color):
    """Calculate relative luminance of a color (0-1 scale, 0=dark, 1=light)."""
    try:
        r, g, b = hex_to_rgb(hex_color)
        # Convert to sRGB
        r, g, b = r / 255.0, g / 255.0, b / 255.0
        # Apply gamma correction
        r = r / 12.92 if r <= 0.03928 else ((r + 0.055) / 1.055) ** 2.4
        g = g / 12.92 if g <= 0.03928 else ((g + 0.055) / 1.055) ** 2.4
        b = b / 12.92 if b <= 0.03928 else ((b + 0.055) / 1.055) ** 2.4
        # Calculate luminance
        return 0.2126 * r + 0.7152 * g + 0.0722 * b
    except:
        return 0.5  # Default to mid-range if parsing fails

def is_light_color(hex_color):
    """Determine if a color is light (True) or dark (False)."""
    return get_luminance(hex_color) > 0.5

def get_contrast_text_color(bg_color):
    """Return appropriate text color (light or dark) based on background."""
    if is_light_color(bg_color):
        return "#0f172a"  # Dark text for light backgrounds
    else:
        return "#f8fafc"  # Light text for dark backgrounds

def get_contrast_button_text(button_bg_color):
    """Return appropriate button text color based on button background."""
    if is_light_color(button_bg_color):
        return "#0f172a"  # Dark text for light button backgrounds
    else:
        return "#ffffff"  # White text for dark button backgrounds

def ensure_color_contrast(design):
    """Ensure all colors in a design have proper contrast for readability."""
    bg_color = design.get('backgroundColor', '#ffffff')
    accent_color = design.get('accentColor', '#3b82f6')

    # Set text color based on background luminance
    design['textColor'] = get_contrast_text_color(bg_color)

    # Set button/accent text color based on accent luminance
    design['buttonTextColor'] = get_contrast_button_text(accent_color)

    # Add secondary text color (slightly dimmed)
    if is_light_color(bg_color):
        design['secondaryTextColor'] = "#4b5563"  # Gray for light bg
    else:
        design['secondaryTextColor'] = "#9ca3af"  # Light gray for dark bg

    # For vibrant/multicolor card designs, ensure card text is readable
    if design.get('multiColorCards'):
        design['cardTextColor'] = "#ffffff"  # Always white on gradient cards

    # Ensure accent hover color contrasts well
    if is_light_color(accent_color):
        design['accentHoverColor'] = "#0f172a"  # Dark hover for light accent
    else:
        design['accentHoverColor'] = "#f8fafc"  # Light hover for dark accent

    return design

# DISTINCT DESIGN CONCEPTS - Each represents a unique visual philosophy
# These are complete design systems, not just color variations

DESIGN_CONCEPTS = [
    # 1. Glassmorphism - Frosted glass with blur effects
    {
        "conceptId": "glassmorphism",
        "name": "Glass Morphism",
        "description": "Frosted glass cards with backdrop blur",
        "layoutType": "cards",
        "backgroundColor": "#0f172a",
        "textColor": "#f1f5f9",
        "accentColor": "#8b5cf6",
        "cardBg": "rgba(255, 255, 255, 0.05)",
        "cardBorder": "1px solid rgba(255, 255, 255, 0.1)",
        "cardShadow": "0 8px 32px rgba(0, 0, 0, 0.3)",
        "backdropBlur": "blur(12px)",
        "fontFamily": "Inter",
        "borderRadius": 20,
        "spacing": "relaxed",
        "iconStyle": "outlined",
        "animationStyle": "fade-up",
    },
    # 2. Neumorphism - Soft extruded shapes
    {
        "conceptId": "neumorphism",
        "name": "Soft UI",
        "description": "Soft shadows creating 3D depth",
        "layoutType": "grid",
        "backgroundColor": "#e0e5ec",
        "textColor": "#2d3436",
        "accentColor": "#6c5ce7",
        "cardBg": "#e0e5ec",
        "cardBorder": "none",
        "cardShadow": "8px 8px 16px #c8ccd4, -8px -8px 16px #ffffff",
        "fontFamily": "Poppins",
        "borderRadius": 24,
        "spacing": "normal",
        "iconStyle": "filled",
        "animationStyle": "scale",
    },
    # 3. Brutalist - Raw, bold, unconventional
    {
        "conceptId": "brutalist",
        "name": "Brutalist",
        "description": "Bold borders, raw typography",
        "layoutType": "stacked",
        "backgroundColor": "#ffffff",
        "textColor": "#000000",
        "accentColor": "#ff0000",
        "cardBg": "#ffffff",
        "cardBorder": "4px solid #000000",
        "cardShadow": "8px 8px 0px #000000",
        "fontFamily": "Space Grotesk",
        "borderRadius": 0,
        "spacing": "compact",
        "iconStyle": "bold",
        "animationStyle": "none",
    },
    # 4. Gradient Mesh - Colorful gradients
    {
        "conceptId": "gradient-mesh",
        "name": "Aurora Gradient",
        "description": "Flowing gradient backgrounds",
        "layoutType": "timeline",
        "backgroundColor": "#0c0a1d",
        "textColor": "#ffffff",
        "accentColor": "#f472b6",
        "cardBg": "linear-gradient(135deg, rgba(139, 92, 246, 0.2) 0%, rgba(244, 114, 182, 0.2) 100%)",
        "cardBorder": "1px solid rgba(255, 255, 255, 0.1)",
        "cardShadow": "0 4px 24px rgba(139, 92, 246, 0.3)",
        "gradient": "linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f472b6 100%)",
        "fontFamily": "DM Sans",
        "borderRadius": 16,
        "spacing": "relaxed",
        "iconStyle": "gradient",
        "animationStyle": "float",
        "glowEffect": True,
    },
    # 5. Minimal Line Art - Clean lines, minimal colors
    {
        "conceptId": "minimal-lines",
        "name": "Line Minimal",
        "description": "Clean lines, subtle elegance",
        "layoutType": "list",
        "backgroundColor": "#fafafa",
        "textColor": "#171717",
        "accentColor": "#0ea5e9",
        "cardBg": "transparent",
        "cardBorder": "1px solid #e5e5e5",
        "cardShadow": "none",
        "fontFamily": "Inter",
        "borderRadius": 8,
        "spacing": "compact",
        "iconStyle": "line",
        "animationStyle": "slide",
        "dividerStyle": "thin-line",
    },
    # 6. Neon Glow - Cyberpunk aesthetic
    {
        "conceptId": "neon-glow",
        "name": "Cyber Neon",
        "description": "Glowing neon accents on dark",
        "layoutType": "cards",
        "backgroundColor": "#0a0a0f",
        "textColor": "#e0e0e0",
        "accentColor": "#00ff88",
        "cardBg": "rgba(0, 255, 136, 0.05)",
        "cardBorder": "1px solid #00ff88",
        "cardShadow": "0 0 20px rgba(0, 255, 136, 0.3), inset 0 0 20px rgba(0, 255, 136, 0.05)",
        "fontFamily": "JetBrains Mono",
        "borderRadius": 12,
        "spacing": "normal",
        "iconStyle": "neon",
        "animationStyle": "glow-pulse",
        "glowEffect": True,
    },
    # 7. Paper Cutout - Layered paper effect
    {
        "conceptId": "paper-cutout",
        "name": "Paper Layers",
        "description": "Stacked paper with subtle shadows",
        "layoutType": "alternating",
        "backgroundColor": "#f5f5f4",
        "textColor": "#292524",
        "accentColor": "#ea580c",
        "cardBg": "#ffffff",
        "cardBorder": "none",
        "cardShadow": "0 1px 3px rgba(0, 0, 0, 0.08), 0 4px 12px rgba(0, 0, 0, 0.05)",
        "fontFamily": "Merriweather",
        "borderRadius": 4,
        "spacing": "relaxed",
        "iconStyle": "minimal",
        "animationStyle": "lift",
        "layeredEffect": True,
    },
    # 8. Duotone - Two-color harmony
    {
        "conceptId": "duotone",
        "name": "Duotone Harmony",
        "description": "Two-color elegant design",
        "layoutType": "grid",
        "backgroundColor": "#1e1b4b",
        "textColor": "#e0e7ff",
        "accentColor": "#fbbf24",
        "secondaryAccent": "#a78bfa",
        "cardBg": "rgba(167, 139, 250, 0.1)",
        "cardBorder": "1px solid rgba(167, 139, 250, 0.2)",
        "cardShadow": "0 4px 16px rgba(0, 0, 0, 0.2)",
        "fontFamily": "Quicksand",
        "borderRadius": 16,
        "spacing": "normal",
        "iconStyle": "duotone",
        "animationStyle": "fade",
    },
    # 9. Retro Terminal - Old-school computer aesthetic
    {
        "conceptId": "retro-terminal",
        "name": "Retro Terminal",
        "description": "Classic terminal with scan lines",
        "layoutType": "list",
        "backgroundColor": "#0d1117",
        "textColor": "#39d353",
        "accentColor": "#58a6ff",
        "cardBg": "rgba(57, 211, 83, 0.05)",
        "cardBorder": "1px dashed #39d353",
        "cardShadow": "none",
        "fontFamily": "JetBrains Mono",
        "borderRadius": 0,
        "spacing": "compact",
        "iconStyle": "ascii",
        "animationStyle": "typewriter",
        "scanLines": True,
    },
    # 10. Luxury Gold - Premium elegant
    {
        "conceptId": "luxury-gold",
        "name": "Luxury Gold",
        "description": "Premium dark with gold accents",
        "layoutType": "timeline",
        "backgroundColor": "#18181b",
        "textColor": "#fafafa",
        "accentColor": "#fbbf24",
        "cardBg": "linear-gradient(135deg, rgba(251, 191, 36, 0.08) 0%, rgba(251, 191, 36, 0.02) 100%)",
        "cardBorder": "1px solid rgba(251, 191, 36, 0.3)",
        "cardShadow": "0 4px 24px rgba(0, 0, 0, 0.4)",
        "fontFamily": "Playfair Display",
        "borderRadius": 8,
        "spacing": "relaxed",
        "iconStyle": "elegant",
        "animationStyle": "shine",
    },
    # 11. Organic Shapes - Blob-like, natural forms
    {
        "conceptId": "organic-shapes",
        "name": "Organic Flow",
        "description": "Natural blob shapes, soft curves",
        "layoutType": "floating",
        "backgroundColor": "#fef3c7",
        "textColor": "#78350f",
        "accentColor": "#f97316",
        "cardBg": "#ffffff",
        "cardBorder": "none",
        "cardShadow": "0 10px 40px rgba(249, 115, 22, 0.15)",
        "fontFamily": "Nunito",
        "borderRadius": 32,
        "spacing": "relaxed",
        "iconStyle": "rounded",
        "animationStyle": "morph",
        "blobBackground": True,
    },
    # 12. Monochrome Elegance
    {
        "conceptId": "monochrome",
        "name": "Mono Elegance",
        "description": "Single color, multiple shades",
        "layoutType": "minimal",
        "backgroundColor": "#ffffff",
        "textColor": "#18181b",
        "accentColor": "#18181b",
        "cardBg": "#fafafa",
        "cardBorder": "1px solid #e5e5e5",
        "cardShadow": "0 2px 8px rgba(0, 0, 0, 0.04)",
        "fontFamily": "Inter",
        "borderRadius": 12,
        "spacing": "normal",
        "iconStyle": "solid",
        "animationStyle": "subtle",
    },
    # 13. Vibrant Cards - Bold, colorful cards
    {
        "conceptId": "vibrant-cards",
        "name": "Vibrant Pop",
        "description": "Bold colorful card backgrounds",
        "layoutType": "grid",
        "backgroundColor": "#fafafa",
        "textColor": "#ffffff",
        "accentColor": "#8b5cf6",
        "cardBg": "var(--card-gradient)",  # Each card gets unique gradient
        "cardBorder": "none",
        "cardShadow": "0 8px 24px rgba(0, 0, 0, 0.12)",
        "fontFamily": "Poppins",
        "borderRadius": 20,
        "spacing": "normal",
        "iconStyle": "white",
        "animationStyle": "bounce",
        "multiColorCards": True,
        "cardGradients": [
            "linear-gradient(135deg, #667eea 0%, #764ba2 100%)",
            "linear-gradient(135deg, #f093fb 0%, #f5576c 100%)",
            "linear-gradient(135deg, #4facfe 0%, #00f2fe 100%)",
            "linear-gradient(135deg, #43e97b 0%, #38f9d7 100%)",
        ],
    },
    # 14. Dark Forest - Nature-inspired dark
    {
        "conceptId": "dark-forest",
        "name": "Dark Forest",
        "description": "Deep greens with organic feel",
        "layoutType": "journey",
        "backgroundColor": "#14532d",
        "textColor": "#dcfce7",
        "accentColor": "#4ade80",
        "cardBg": "rgba(74, 222, 128, 0.08)",
        "cardBorder": "1px solid rgba(74, 222, 128, 0.2)",
        "cardShadow": "0 4px 20px rgba(0, 0, 0, 0.3)",
        "fontFamily": "DM Sans",
        "borderRadius": 16,
        "spacing": "relaxed",
        "iconStyle": "nature",
        "animationStyle": "grow",
    },
    # 15. Ocean Depths - Deep blue oceanic
    {
        "conceptId": "ocean-depths",
        "name": "Ocean Depths",
        "description": "Deep sea blues with wave motifs",
        "layoutType": "horizontal",
        "backgroundColor": "#0c1929",
        "textColor": "#e0f2fe",
        "accentColor": "#22d3ee",
        "cardBg": "rgba(34, 211, 238, 0.08)",
        "cardBorder": "1px solid rgba(34, 211, 238, 0.15)",
        "cardShadow": "0 4px 24px rgba(34, 211, 238, 0.2)",
        "fontFamily": "Space Grotesk",
        "borderRadius": 16,
        "spacing": "normal",
        "iconStyle": "wave",
        "animationStyle": "wave",
    },
]

# Layout types available for each context
LAYOUT_TYPES = {
    "education": ["cards", "timeline", "alternating", "journey", "floating", "grid", "list", "minimal"],
    "career": ["timeline", "cards", "stacked", "horizontal", "list", "grid", "minimal"],
    "profile": ["horizontal", "centered", "split", "minimal", "creative", "cards"],
    "awards": ["cards", "grid", "list", "timeline", "showcase"],
    "general": ["cards", "grid", "list", "timeline", "stacked"],
}

FONT_FAMILIES = [
    "Inter", "Poppins", "Space Grotesk", "JetBrains Mono", "DM Sans",
    "Playfair Display", "Merriweather", "Quicksand", "IBM Plex Sans",
    "Cormorant Garamond", "Nunito", "Roboto", "Lato", "Montserrat"
]

@api_view(["POST"])
@permission_classes([IsAuthenticated])
def generate_designs(request):
    """
    Generate truly distinct design concepts - each with unique visual philosophy,
    not just color variations. Ensures proper color contrast for readability.
    """
    context = request.data.get("context", "general")
    count = request.data.get("count", 5)

    # Get available layouts for this context
    available_layouts = LAYOUT_TYPES.get(context.lower(), LAYOUT_TYPES["general"])

    # Shuffle design concepts to get random selection
    shuffled_concepts = DESIGN_CONCEPTS.copy()
    random.shuffle(shuffled_concepts)

    # Select unique concepts
    selected_concepts = shuffled_concepts[:count]

    print(f"Generating {count} distinct design concepts for context: {context}")

    generated_designs = []
    used_layouts = set()

    for i, concept in enumerate(selected_concepts):
        # Deep copy the concept
        design = concept.copy()

        # Assign unique ID
        design["id"] = f"{concept['conceptId']}-{random.randint(1000, 9999)}"

        # Try to use a unique layout for each design
        preferred_layout = concept.get("layoutType", "cards")
        if preferred_layout in available_layouts and preferred_layout not in used_layouts:
            design["layoutType"] = preferred_layout
        else:
            # Find an unused layout
            unused_layouts = [l for l in available_layouts if l not in used_layouts]
            if unused_layouts:
                design["layoutType"] = random.choice(unused_layouts)
            else:
                design["layoutType"] = random.choice(available_layouts)

        used_layouts.add(design["layoutType"])

        # Ensure color contrast
        design = ensure_color_contrast(design)

        # Add context-specific properties
        if "profile" in context.lower():
            design["imageStyle"] = random.choice(["circle", "rounded", "square"])

        generated_designs.append(design)

    # If we need more designs than available concepts, create variations
    while len(generated_designs) < count:
        # Pick a random concept and create a variation
        base_concept = random.choice(DESIGN_CONCEPTS).copy()
        variation = base_concept.copy()
        variation["id"] = f"{base_concept['conceptId']}-var-{random.randint(1000, 9999)}"
        variation["name"] = f"{base_concept['name']} Alt"

        # Use a different layout
        unused_layouts = [l for l in available_layouts if l not in used_layouts]
        if unused_layouts:
            variation["layoutType"] = random.choice(unused_layouts)
            used_layouts.add(variation["layoutType"])
        else:
            variation["layoutType"] = random.choice(available_layouts)

        # Slightly vary some properties
        variation["borderRadius"] = random.choice([4, 8, 12, 16, 20, 24, 32])
        variation["spacing"] = random.choice(["compact", "normal", "relaxed"])
        variation["fontFamily"] = random.choice(FONT_FAMILIES)

        variation = ensure_color_contrast(variation)
        generated_designs.append(variation)

    print(f"Generated {len(generated_designs)} distinct designs with proper color contrast")
    return Response({
        "designs": generated_designs[:count],
        "source": "distinct_concepts"
    })


@api_view(["POST"])
@permission_classes([IsAuthenticated])
def import_from_cv(request):
    """
    Parse the CV uploaded in user's profile and import education/career/skills data.
    """
    from PyPDF2 import PdfReader
    from docx import Document
    from profiles.models import UserProfile
    from education.models import EducationEntry
    from career.models import CareerEntry, Skill
    from django.shortcuts import get_object_or_404

    user = request.user
    section = request.data.get("section", "education")  # education, career, skills, or all

    # Auto-create profile if it doesn't exist
    profile, created = UserProfile.objects.get_or_create(user=user)

    if not profile.cv:
        return Response({'error': 'No CV uploaded. Please upload your CV first using the onboarding page or Profile section.'},
                       status=status.HTTP_400_BAD_REQUEST)

    try:
        # Get the file path and read the CV
        cv_path = profile.cv.path
        filename = cv_path.lower()
        extracted_text = ""

        # Extract text based on file type
        if filename.endswith('.pdf'):
            with open(cv_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
        elif filename.endswith('.docx'):
            doc = Document(cv_path)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
        else:
            return Response({'error': 'Unsupported file format. CV must be PDF or DOCX.'},
                          status=status.HTTP_400_BAD_REQUEST)

        if not extracted_text.strip():
            return Response({'error': 'Could not extract text from the CV.'},
                          status=status.HTTP_400_BAD_REQUEST)

        result = {'success': True, 'imported': {}}

        # Parse Education
        if section in ["education", "all"]:
            edu_prompt = f"""Parse this resume/CV and extract ONLY the education information.
Return a JSON array with the following structure:

[
    {{
        "degree": "Degree Name (e.g., Bachelor of Science in Computer Science)",
        "university": "Institution/University Name",
        "year": "Year or Year Range (e.g., 2015-2019 or 2019)",
        "description": "Any additional details, achievements, or GPA"
    }}
]

RESUME TEXT:
{extracted_text[:8000]}

Return ONLY the JSON array, no explanation or markdown. If no education found, return an empty array []."""

            response = model.generate_content(edu_prompt)
            response_text = response.text.strip()
            response_text = re.sub(r'^```(?:json)?\s*', '', response_text)
            response_text = re.sub(r'\s*```$', '', response_text)

            education_data = json.loads(response_text)

            imported_edu = 0
            for edu in education_data:
                if edu.get('degree') or edu.get('university'):
                    EducationEntry.objects.create(
                        user=user,
                        degree=edu.get('degree', ''),
                        university=edu.get('university', ''),
                        year=edu.get('year', ''),
                        description=edu.get('description', '')
                    )
                    imported_edu += 1
            result['imported']['education'] = imported_edu

        # Parse Career/Experience
        if section in ["career", "all"]:
            career_prompt = f"""Parse this resume/CV and extract ONLY the work experience/career information.
Return a JSON array with the following structure:

[
    {{
        "title": "Job Title",
        "company": "Company Name",
        "year": "Year range (e.g., 2020-2022 or 2020-Present)",
        "description": "Job description and achievements"
    }}
]

RESUME TEXT:
{extracted_text[:8000]}

Return ONLY the JSON array, no explanation or markdown. If no experience found, return an empty array []."""

            response = model.generate_content(career_prompt)
            response_text = response.text.strip()
            response_text = re.sub(r'^```(?:json)?\s*', '', response_text)
            response_text = re.sub(r'\s*```$', '', response_text)

            career_data = json.loads(response_text)

            imported_career = 0
            for exp in career_data:
                if exp.get('title') or exp.get('company'):
                    CareerEntry.objects.create(
                        user=user,
                        title=exp.get('title', ''),
                        company=exp.get('company', ''),
                        year=exp.get('year', ''),
                        description=exp.get('description', '')
                    )
                    imported_career += 1
            result['imported']['career'] = imported_career

        # Parse Skills
        if section in ["skills", "all"]:
            skills_prompt = f"""Parse this resume/CV and extract ONLY the skills.
Return a JSON array of skill names:

["skill1", "skill2", "skill3"]

RESUME TEXT:
{extracted_text[:8000]}

Return ONLY the JSON array of strings, no explanation or markdown. If no skills found, return an empty array []."""

            response = model.generate_content(skills_prompt)
            response_text = response.text.strip()
            response_text = re.sub(r'^```(?:json)?\s*', '', response_text)
            response_text = re.sub(r'\s*```$', '', response_text)

            skills_data = json.loads(response_text)

            imported_skills = 0
            for skill_name in skills_data:
                if skill_name and isinstance(skill_name, str):
                    Skill.objects.get_or_create(user=user, name=skill_name.strip())
                    imported_skills += 1
            result['imported']['skills'] = imported_skills

        total = sum(result['imported'].values())
        result['message'] = f'Successfully imported {total} items from your CV.'
        return Response(result)

    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error: {e}")
        return Response({'error': f'Failed to parse AI response. Please try again.'},
                       status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        logger.error(f"CV import error: {e}")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
